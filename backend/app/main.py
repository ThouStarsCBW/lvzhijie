from __future__ import annotations

import json
import os
import re
from io import BytesIO
from html import unescape
from pathlib import Path
from typing import Any
from urllib.error import URLError
from urllib.request import Request, urlopen

from difflib import SequenceMatcher
from docx import Document as DocxDocument
from fastapi import FastAPI, File, Form, HTTPException, Query, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from fastapi.staticfiles import StaticFiles

from app.diffing import build_char_diff, build_paragraph_diff, summarize_legal_risks
from app.legal_search_adapter import DelilegalSearchClient, LegalSearchApiError
from app.models import (
    ActivityEvent,
    AgentArchitecture,
    AgentDepartment,
    AgentGroup,
    BindConversationCaseRequest,
    BranchCreateRequest,
    BranchRevisionCreateRequest,
    Case,
    CaseCreateRequest,
    CaseMemory,
    CaseTask,
    CaseTaskComment,
    CreateCaseFromConversationRequest,
    DocumentCreateRequest,
    FollowUpQuestion,
    LegalAgent,
    LegalDocument,
    LegalDocumentAnalysis,
    LegalDocumentBranch,
    LegalDocumentDiff,
    LegalDocumentRevision,
    LegalResearchResult,
    LegalResearchRun,
    LegalReplyJob,
    LegalReasoningRun,
    MemoryCreateRequest,
    MockConversationCreateRequest,
    MockConversationUpdateRequest,
    OpenClawConnection,
    ReasoningEdge,
    ReasoningNode,
    RevisionCreateRequest,
    ReplyJobCreateRequest,
    SendFollowUpRequest,
    SendMessageRequest,
    TaskCommentCreateRequest,
    TaskCreateRequest,
    TaskExecuteRequest,
    TaskUpdateRequest,
    WechatContact,
    WechatConversation,
    WechatMessage,
    now_iso,
)
from app.mock_wechat_store import MockWechatStore
from app.openclaw_adapter import OpenClawWechatAdapter
from app.store import JsonStore


def load_local_env() -> None:
    env_paths = [
        Path(__file__).resolve().parents[1] / ".env",
        Path(__file__).resolve().parents[2] / ".env",
    ]
    for env_path in env_paths:
        if not env_path.exists():
            continue
        for raw_line in env_path.read_text(encoding="utf-8").splitlines():
            line = raw_line.strip()
            if not line or line.startswith("#") or "=" not in line:
                continue
            key, value = line.split("=", 1)
            key = key.strip()
            value = value.strip().strip('"').strip("'")
            if key:
                os.environ.setdefault(key, value)


load_local_env()

app = FastAPI(title="Lvzhijie Legal Workspace API", version="0.1.0")
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:5173", "http://127.0.0.1:5173"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

store = JsonStore(Path(__file__).parent / "data" / "store.json")
mock_wechat = MockWechatStore(Path(__file__).parent / "data" / "mock_wechat")

app.mount(
    "/mock-wechat-assets",
    StaticFiles(directory=Path(__file__).parent / "data" / "mock_wechat" / "assets"),
    name="mock-wechat-assets",
)


def connection() -> OpenClawConnection:
    return OpenClawConnection.model_validate(store.data["openclaw_connection"])


def using_mock_wechat() -> bool:
    return connection().transport_mode == "mock"


def sync_mock_wechat_if_needed() -> None:
    if using_mock_wechat():
        mock_wechat.sync_to_json_store(store)


def record(event_type: str, title: str, description: str = "", **refs: str | None) -> None:
    store.add(
        "activity_events",
        ActivityEvent(
            event_type=event_type,
            title=title,
            description=description,
            entity_type=refs.get("entity_type"),
            entity_id=refs.get("entity_id"),
        ),
    )


def delete_document_rows(document_id: str) -> None:
    store.remove_where("legal_document_revisions", lambda row: row.get("document_id") == document_id)
    store.remove_where("legal_document_diffs", lambda row: row.get("document_id") == document_id)
    store.remove_where("legal_document_branches", lambda row: row.get("document_id") == document_id)
    store.remove_where("legal_document_analyses", lambda row: row.get("document_id") == document_id)
    store.delete("legal_documents", document_id)


def make_revision_short_hash(revision: LegalDocumentRevision) -> str:
    return revision.id.replace("rev_", "")[:7]


def create_default_branch_for_document(document: LegalDocument, first_revision: LegalDocumentRevision) -> LegalDocumentBranch:
    branch = LegalDocumentBranch(
        document_id=document.id,
        name="main",
        head_revision_id=first_revision.id,
        base_revision_id=first_revision.id,
        is_default=True,
    )
    first_revision.branch_id = branch.id
    first_revision.parent_revision_id = None
    first_revision.created_from_revision_id = None
    first_revision.short_hash = make_revision_short_hash(first_revision)
    document.default_branch_id = branch.id
    document.current_revision_id = first_revision.id
    store.add("legal_document_branches", branch)
    return branch


def document_branches(document_id: str) -> list[LegalDocumentBranch]:
    return sorted(
        store.filter("legal_document_branches", LegalDocumentBranch, document_id=document_id),
        key=lambda item: (not item.is_default, item.name.lower(), item.created_at),
    )


def get_document_branch(document_id: str, branch_id: str) -> LegalDocumentBranch:
    branch = store.get("legal_document_branches", branch_id, LegalDocumentBranch)
    if not branch or branch.document_id != document_id:
        raise HTTPException(status_code=404, detail="Branch not found")
    return branch


def create_revision_on_branch(
    *,
    document: LegalDocument,
    branch: LegalDocumentBranch,
    content_text: str,
    source_filename: str | None,
    author_type: str,
    change_summary: str,
) -> LegalDocumentRevision:
    revisions = store.filter("legal_document_revisions", LegalDocumentRevision, document_id=document.id)
    parent_id = branch.head_revision_id
    revision = LegalDocumentRevision(
        document_id=document.id,
        version_number=len(revisions) + 1,
        content_text=content_text,
        source_filename=source_filename,
        author_type=author_type,
        change_summary=change_summary,
        branch_id=branch.id,
        parent_revision_id=parent_id,
        created_from_revision_id=parent_id,
    )
    revision.short_hash = make_revision_short_hash(revision)
    branch.head_revision_id = revision.id
    branch.updated_at = now_iso()
    document.updated_at = now_iso()
    if branch.is_default:
        document.current_revision_id = revision.id
    store.add("legal_document_revisions", revision)
    store.update("legal_document_branches", branch)
    store.update("legal_documents", document)
    return revision


def delete_wechat_conversation_rows(conversation: WechatConversation) -> dict[str, object]:
    removed_messages = store.remove_where(
        "wechat_messages",
        lambda row: row.get("conversation_id") == conversation.id,
    )
    for case in store.list("cases", Case):
        if case.conversation_ref == conversation.id or case.id == conversation.case_id:
            case.conversation_ref = None
            case.wechat_contact_ref = None
            case.updated_at = now_iso()
            store.update("cases", case)
    deleted = store.delete("wechat_conversations", conversation.id)
    remaining_contact_refs = {
        item.get("contact_id")
        for item in store.data.get("wechat_conversations", [])
        if isinstance(item, dict)
    }
    if conversation.contact_id not in remaining_contact_refs:
        store.delete("wechat_contacts", conversation.contact_id)
    return {"deleted": deleted, "messages": removed_messages}


async def read_legal_file_text(upload: UploadFile) -> str:
    filename = upload.filename or "upload.txt"
    suffix = Path(filename).suffix.lower()
    content = await upload.read()
    if suffix == ".docx":
        document = DocxDocument(BytesIO(content))
        parts = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts)
    if suffix in {".txt", ".md", ""}:
        for encoding in ("utf-8-sig", "utf-8", "gb18030"):
            try:
                return content.decode(encoding)
            except UnicodeDecodeError:
                continue
        return content.decode("utf-8", errors="replace")
    raise HTTPException(status_code=415, detail="Only .txt, .md and .docx files are supported")


CASE_TYPE_LABELS = {
    "contract": "合同纠纷",
    "labor": "劳动争议",
    "marriage": "婚姻家事",
    "debt": "债权债务",
    "traffic": "交通事故",
    "company": "公司商事",
    "real_estate": "房产纠纷",
    "criminal": "刑事法律事务",
    "other": "综合法律咨询",
}

CASE_RULE_HINTS = {
    "contract": "重点核对合同成立、生效、履行、违约、解除和争议解决条款。",
    "labor": "重点核对劳动关系、工资标准、欠薪期间、解除事实、仲裁时效和证据材料。",
    "marriage": "重点核对身份关系、共同财产、债务来源、子女抚养和财产处分证据。",
    "debt": "重点核对借贷合意、款项交付、还款约定、催收记录和诉讼时效。",
    "traffic": "重点核对事故责任、伤情损失、保险信息、医疗票据和误工证明。",
    "company": "重点核对公司决议、股东权利、交易文件、履行记录和内部审批。",
    "real_estate": "重点核对合同、产权状态、付款节点、交付验收和违约责任。",
    "criminal": "重点核对涉嫌罪名、行为事实、主观状态、证据来源和程序节点。",
    "other": "重点核对法律关系、请求基础、证据来源、时效和可执行路径。",
}

CASE_TYPE_VALUES = set(CASE_TYPE_LABELS)

CASE_TYPE_KEYWORDS = {
    "labor": ["工资", "劳动", "入职", "离职", "公司拖欠", "仲裁", "社保", "工伤", "解除"],
    "contract": ["合同", "违约", "履行", "验收", "定金", "条款", "付款期限"],
    "debt": ["借款", "欠款", "还钱", "转账", "债务", "借条", "催收"],
    "marriage": ["离婚", "抚养", "夫妻", "婚姻", "财产分割", "继承"],
    "traffic": ["交通事故", "追尾", "保险", "交警", "责任认定", "伤残"],
    "company": ["股东", "公司", "股权", "合伙", "章程", "分红", "决议"],
    "real_estate": ["房屋", "房产", "买房", "租房", "物业", "交房", "产权"],
    "criminal": ["刑事", "拘留", "取保", "立案", "诈骗", "寻衅滋事", "派出所"],
}

CORE_DEPARTMENTS = [
    AgentDepartment(
        id="criminal_department",
        title="刑事法律事务部",
        description="处理刑事咨询、控告申诉、取保候审与辩护策略。",
        case_types=["criminal"],
    ),
    AgentDepartment(
        id="civil_commercial_department",
        title="民商法律事务部",
        description="处理合同、债权债务、公司商事和一般民事争议。",
        case_types=["contract", "debt", "company"],
    ),
    AgentDepartment(
        id="traffic_department",
        title="交通法律事务部",
        description="处理交通事故责任、赔偿项目、保险理赔和调解诉讼。",
        case_types=["traffic"],
    ),
    AgentDepartment(
        id="ip_department",
        title="知识产权法律事务部",
        description="处理著作权、商标、商业秘密和知识产权合同风险。",
        case_types=["other"],
    ),
    AgentDepartment(
        id="marriage_wealth_department",
        title="婚姻家事与财富传承法律事务部",
        description="处理婚姻家事、继承、财富规划和家族财产安排。",
        case_types=["marriage"],
    ),
    AgentDepartment(
        id="tax_department",
        title="财税法律事务部",
        description="处理财税合规、交易税负、发票风险和税务争议。",
        case_types=["company", "other"],
    ),
    AgentDepartment(
        id="labor_department",
        title="劳动法律事务部",
        description="处理劳动合同、欠薪、解除、工伤和劳动仲裁。",
        case_types=["labor"],
    ),
]

CLIENT_SERVICE_DEPARTMENTS = [
    AgentDepartment(
        id="legal_consultation_service",
        title="法律咨询",
        description="承接初步法律咨询、问题识别、沟通口径和下一步材料清单。",
        case_types=["other"],
    ),
    AgentDepartment(
        id="intake_quote_service",
        title="接案报价",
        description="根据案件类型、工作量、风险和交付范围形成接案与报价建议。",
        case_types=["contract", "labor", "marriage", "debt", "traffic", "company", "real_estate", "criminal", "other"],
    ),
    AgentDepartment(
        id="complaint_handling_service",
        title="投诉处理",
        description="记录客户投诉、识别服务问题、推动内部处理和反馈闭环。",
        case_types=["other"],
    ),
]

COMPLIANCE_REVIEW_DEPARTMENTS = [
    AgentDepartment(
        id="case_acceptance_approval",
        title="收案审批",
        description="审查案件是否满足收案条件、材料基础、服务边界和人工复核要求。",
        case_types=["contract", "labor", "marriage", "debt", "traffic", "company", "real_estate", "criminal", "other"],
    ),
    AgentDepartment(
        id="conflict_check",
        title="利益冲突审查",
        description="核查客户、相对方、关联主体和既有案件之间的潜在利益冲突。",
        case_types=["contract", "labor", "marriage", "debt", "traffic", "company", "real_estate", "criminal", "other"],
    ),
    AgentDepartment(
        id="risk_assessment",
        title="风险评估",
        description="评估事实不确定性、证据缺口、合规风险、声誉风险和对外表述边界。",
        case_types=["contract", "labor", "marriage", "debt", "traffic", "company", "real_estate", "criminal", "other"],
    ),
]

ARCHIVE_MANAGEMENT_DEPARTMENTS = [
    AgentDepartment(
        id="case_file_archiving",
        title="案卷归档",
        description="归档案件材料、版本文书、证据目录、推理记录和过程留痕。",
        case_types=["contract", "labor", "marriage", "debt", "traffic", "company", "real_estate", "criminal", "other"],
    ),
    AgentDepartment(
        id="archive_search",
        title="档案查询",
        description="按案件、客户、文件类型、时间线和证据标签查询历史档案。",
        case_types=["contract", "labor", "marriage", "debt", "traffic", "company", "real_estate", "criminal", "other"],
    ),
]


def case_type_label(case_type: str) -> str:
    return CASE_TYPE_LABELS.get(case_type, "综合法律咨询")


def case_rule_hint(case_type: str) -> str:
    return CASE_RULE_HINTS.get(case_type, CASE_RULE_HINTS["other"])


def first_sentence(value: str, fallback: str = "暂未形成摘要。") -> str:
    text = " ".join(value.split())
    if not text:
        return fallback
    for marker in ("。", "；", ";", "."):
        if marker in text:
            return text.split(marker, 1)[0] + marker
    return text[:120]


def message_plaintext(message: WechatMessage) -> str:
    attachment_names = "、".join(attachment.name for attachment in message.attachments)
    content = message.content.strip()
    if attachment_names:
        return f"{content} [附件：{attachment_names}]" if content else f"[附件：{attachment_names}]"
    return content


def build_conversation_transcript(messages: list[WechatMessage]) -> str:
    sender_labels = {
        "wechat_user": "客户",
        "openclaw_auto": "微信桥自动回复",
        "owner": "我方",
        "system": "系统",
    }
    lines = []
    for message in messages:
        text = message_plaintext(message).strip()
        if not text:
            continue
        sender = sender_labels.get(message.sender, message.sender)
        lines.append(f"{message.created_at} {sender}：{text}")
    return "\n".join(lines)


def infer_case_type_from_text(text: str) -> str:
    scores: dict[str, int] = {}
    for case_type, keywords in CASE_TYPE_KEYWORDS.items():
        scores[case_type] = sum(1 for keyword in keywords if keyword in text)
    best_type, best_score = max(scores.items(), key=lambda item: item[1])
    return best_type if best_score > 0 else "other"


def fallback_case_analysis(
    *,
    contact: WechatContact | None,
    messages: list[WechatMessage],
    requested_title: str | None = None,
    requested_case_type: str | None = None,
) -> dict[str, object]:
    message_texts = [message_plaintext(message) for message in messages if message_plaintext(message).strip()]
    inbound_texts = [
        message_plaintext(message)
        for message in messages
        if message.direction == "inbound" and message_plaintext(message).strip()
    ]
    transcript_text = " ".join(message_texts)
    inferred_type = requested_case_type if requested_case_type in CASE_TYPE_VALUES else infer_case_type_from_text(transcript_text)
    display_name = contact.display_name if contact else "微信用户"
    title = requested_title or f"{display_name}{case_type_label(inferred_type)}咨询"
    summary_source = " ".join(inbound_texts or message_texts)
    facts = [text[:160] for text in inbound_texts[:5]]
    evidences = []
    for message in messages:
        for attachment in message.attachments:
            evidences.append(f"客户或我方上传附件：{attachment.name}")
    uncertainties = [
        "需进一步确认关键时间节点、证据来源和客户的最终处理目标。",
        f"需按{case_type_label(inferred_type)}方向核对请求基础、时效和可证明材料。",
    ]
    return {
        "title": title[:80],
        "case_type": inferred_type,
        "summary": first_sentence(summary_source, "由微信会话创建，暂未形成充分案情摘要。")[:500],
        "facts": facts or ["已从微信会话创建案件，需继续补充事实。"],
        "evidence": evidences,
        "uncertainties": uncertainties,
        "suggested_tasks": [
            "整理完整时间线",
            "核对证据材料",
            "明确客户诉求和处理路径",
        ],
    }


def extract_json_object(text: str) -> dict[str, Any] | None:
    value = text.strip()
    if not value:
        return None
    try:
        parsed = json.loads(value)
        return parsed if isinstance(parsed, dict) else None
    except json.JSONDecodeError:
        pass
    start = value.find("{")
    end = value.rfind("}")
    if start == -1 or end == -1 or end <= start:
        return None
    try:
        parsed = json.loads(value[start : end + 1])
    except json.JSONDecodeError:
        return None
    return parsed if isinstance(parsed, dict) else None


def normalize_case_analysis(raw: dict[str, Any], fallback: dict[str, object]) -> dict[str, object]:
    case_type = str(raw.get("case_type") or fallback["case_type"])
    if case_type not in CASE_TYPE_VALUES:
        case_type = str(fallback["case_type"])

    def string_list(key: str, fallback_key: str) -> list[str]:
        value = raw.get(key)
        if isinstance(value, list):
            items = [str(item).strip() for item in value if str(item).strip()]
            return items[:8]
        return list(fallback.get(fallback_key, []))[:8]  # type: ignore[arg-type]

    return {
        "title": (str(raw.get("title") or fallback["title"]).strip() or str(fallback["title"]))[:80],
        "case_type": case_type,
        "summary": (str(raw.get("summary") or fallback["summary"]).strip() or str(fallback["summary"]))[:800],
        "facts": string_list("facts", "facts"),
        "evidence": string_list("evidence", "evidence"),
        "uncertainties": string_list("uncertainties", "uncertainties"),
        "suggested_tasks": string_list("suggested_tasks", "suggested_tasks"),
    }


async def analyze_conversation_for_case(
    *,
    contact: WechatContact | None,
    messages: list[WechatMessage],
    requested_title: str | None,
    requested_case_type: str | None,
) -> dict[str, object]:
    fallback = fallback_case_analysis(
        contact=contact,
        messages=messages,
        requested_title=requested_title,
        requested_case_type=requested_case_type,
    )
    api_key = os.getenv("LVZHIJIE_LLM_API_KEY") or os.getenv("OPENAI_API_KEY")
    if not api_key:
        return fallback

    transcript = build_conversation_transcript(messages)
    if not transcript:
        return fallback
    base_url = (os.getenv("LVZHIJIE_LLM_BASE_URL") or os.getenv("OPENAI_BASE_URL") or "https://api.openai.com/v1").rstrip("/")
    model = os.getenv("LVZHIJIE_LLM_MODEL") or os.getenv("OPENAI_MODEL") or "gpt-4o-mini"
    prompt = f"""请根据以下微信咨询聊天记录，生成一份案件建档 JSON。
只能输出 JSON 对象，不要输出 Markdown。
JSON 字段必须包含：
- title: 80字以内案件标题
- case_type: 只能是 contract, labor, marriage, debt, traffic, company, real_estate, criminal, other 之一
- summary: 300字以内案件摘要
- facts: 已掌握事实数组
- evidence: 证据或附件数组
- uncertainties: 仍需追问或核验的不确定点数组
- suggested_tasks: 后续办理任务数组

联系人：{contact.display_name if contact else "微信用户"}
聊天记录：
{transcript[:12000]}
"""
    payload = {
        "model": model,
        "messages": [
            {"role": "system", "content": "你是律所案件接待助手，负责把客户微信咨询整理为结构化案件建档信息。"},
            {"role": "user", "content": prompt},
        ],
        "temperature": 0.2,
    }
    request = Request(
        f"{base_url}/chat/completions",
        data=json.dumps(payload, ensure_ascii=False).encode("utf-8"),
        headers={
            "Authorization": f"Bearer {api_key}",
            "Content-Type": "application/json",
        },
        method="POST",
    )
    try:
        with urlopen(request, timeout=25) as response:
            body = response.read().decode("utf-8")
    except (OSError, URLError, TimeoutError):
        return fallback
    parsed_response = extract_json_object(body)
    choices = parsed_response.get("choices") if parsed_response else None
    content = ""
    if isinstance(choices, list) and choices:
        message = choices[0].get("message") if isinstance(choices[0], dict) else None
        if isinstance(message, dict):
            content = str(message.get("content") or "")
    raw_analysis = extract_json_object(content)
    if raw_analysis is None:
        return fallback
    return normalize_case_analysis(raw_analysis, fallback)


def latest_document_revision(document_id: str) -> LegalDocumentRevision | None:
    revisions = sorted(
        store.filter("legal_document_revisions", LegalDocumentRevision, document_id=document_id),
        key=lambda item: item.version_number,
    )
    return revisions[-1] if revisions else None


def build_docx_stream(title: str, content: str) -> BytesIO:
    document = DocxDocument()
    document.add_heading(title, level=1)
    for block in content.split("\n"):
        line = block.strip()
        if not line:
            continue
        if line.endswith("：") and len(line) <= 24:
            document.add_heading(line[:-1], level=2)
        else:
            document.add_paragraph(line)
    stream = BytesIO()
    document.save(stream)
    stream.seek(0)
    return stream


def create_text_document(
    *,
    case_id: str | None,
    title: str,
    document_type: str,
    content_text: str,
    source_filename: str | None = None,
    change_summary: str = "",
    author_type: str = "agent",
) -> LegalDocument:
    normalized_type = document_type if document_type in {"contract", "letter", "pleading", "evidence", "other"} else "other"
    document = LegalDocument(case_id=case_id, title=title, document_type=normalized_type)  # type: ignore[arg-type]
    revision = LegalDocumentRevision(
        document_id=document.id,
        version_number=1,
        content_text=content_text,
        source_filename=source_filename,
        author_type=author_type,  # type: ignore[arg-type]
        change_summary=change_summary,
    )
    create_default_branch_for_document(document, revision)
    store.add("legal_documents", document)
    store.add("legal_document_revisions", revision)
    return document


def case_memories(case_id: str) -> list[CaseMemory]:
    return sorted(
        store.filter("case_memories", CaseMemory, case_id=case_id),
        key=lambda item: item.created_at,
    )


def case_documents(case_id: str) -> list[LegalDocument]:
    return sorted(
        store.filter("legal_documents", LegalDocument, case_id=case_id),
        key=lambda item: item.updated_at,
        reverse=True,
    )


def case_messages(case: Case) -> list[WechatMessage]:
    if not case.conversation_ref:
        return []
    return sorted(
        store.filter("wechat_messages", WechatMessage, conversation_id=case.conversation_ref),
        key=lambda item: item.created_at,
    )


def infer_case_summary(case: Case, memories: list[CaseMemory], messages: list[WechatMessage]) -> str:
    if case.summary.strip():
        return case.summary.strip()
    if memories:
        return "；".join(memory.content for memory in memories[:3])
    inbound = [message.content for message in messages if message.direction == "inbound"]
    return first_sentence(" ".join(inbound), "该案件仍处于信息收集阶段。")


def build_short_reply(case: Case, payload: ReplyJobCreateRequest, memories: list[CaseMemory]) -> str:
    facts = [memory.content for memory in memories if memory.kind in {"fact", "timeline", "evidence"}][:2]
    missing = [memory.content for memory in memories if memory.kind == "uncertainty"][:2]
    fact_text = "；".join(facts) if facts else "目前信息还不足，需要先补齐关键事实。"
    missing_text = "；".join(missing) if missing else "建议继续补充时间节点、证据材料和您的处理目标。"
    question = payload.user_question.strip() or "您提交的咨询问题"
    return (
        f"关于“{question[:80]}”，我先按{case_type_label(case.case_type)}方向为您做初步梳理："
        f"{fact_text} "
        f"下一步请优先补充：{missing_text} "
        "在材料完整前，当前意见仅作为初步判断，正式处理建议经过人工复核。"
    )


def build_long_reply_content(
    case: Case,
    job: LegalReplyJob,
    memories: list[CaseMemory],
    documents: list[LegalDocument],
    latest_run: LegalReasoningRun | None,
) -> str:
    facts = [memory.content for memory in memories if memory.kind in {"fact", "timeline", "evidence"}]
    uncertainties = [memory.content for memory in memories if memory.kind == "uncertainty"]
    document_lines = []
    for document in documents[:5]:
        revision = latest_document_revision(document.id)
        hint = revision.change_summary if revision else "暂无版本摘要"
        document_lines.append(f"- {document.title}：{hint}")
    followups = latest_run.follow_up_questions if latest_run else []
    facts_text = "\n".join(f"- {item}" for item in facts) or "- 暂无已确认事实，建议先完成事实访谈。"
    uncertainty_text = "\n".join(f"- {item}" for item in uncertainties) or "- 暂未登记重大不确定点。"
    document_text = "\n".join(document_lines) or "- 暂无案件文件。"
    followup_text = "\n".join(f"- {item}" for item in followups[:5]) or "- 请补充关键时间节点、证据材料和处理目标。"
    return f"""法律意见摘要：
案件名称：{case.title}
案件类型：{case_type_label(case.case_type)}
任务标题：{job.title}

一、案件摘要：
{job.case_summary or infer_case_summary(case, memories, case_messages(case))}

二、已掌握事实与证据：
{facts_text}

三、关联文件：
{document_text}

四、初步法律分析：
{case_rule_hint(case.case_type)}
结合现有材料，当前更适合形成阶段性意见，而非直接输出最终结论。需要把事实、证据和诉求目标对应到法律要件后，再决定协商、函件、投诉、仲裁或诉讼路径。

五、风险与不确定点：
{uncertainty_text}

六、建议追问/补充材料：
{followup_text}

七、人工复核提示：
本文件由系统根据案件记录生成，应由律师或办案人员复核事实来源、法律依据、证据完整性和表述边界后再对外使用。
"""


def case_tasks(case_id: str) -> list[CaseTask]:
    return sorted(
        store.filter("case_tasks", CaseTask, case_id=case_id),
        key=lambda item: item.created_at,
        reverse=True,
    )


def task_comments(case_id: str, task_id: str | None = None) -> list[CaseTaskComment]:
    comments = store.filter("case_task_comments", CaseTaskComment, case_id=case_id)
    if task_id:
        comments = [comment for comment in comments if comment.task_id == task_id]
    return sorted(comments, key=lambda item: item.created_at)


def research_runs(case_id: str) -> list[LegalResearchRun]:
    return sorted(
        store.filter("legal_research_runs", LegalResearchRun, case_id=case_id),
        key=lambda item: item.created_at,
        reverse=True,
    )


def research_results(case_id: str, task_id: str | None = None) -> list[LegalResearchResult]:
    results = store.filter("legal_research_results", LegalResearchResult, case_id=case_id)
    if task_id:
        results = [result for result in results if result.task_id == task_id]
    return sorted(results, key=lambda item: item.relevance_score, reverse=True)


def get_case_task(case_id: str, task_id: str) -> CaseTask:
    task = store.get("case_tasks", task_id, CaseTask)
    if not task or task.case_id != case_id:
        raise HTTPException(status_code=404, detail="Task not found")
    return task


def ensure_case_exists(case_id: str) -> Case:
    case = store.get("cases", case_id, Case)
    if not case:
        raise HTTPException(status_code=404, detail="Case not found")
    return case


def add_task_comment(
    *,
    case_id: str,
    task_id: str,
    message: str,
    author_type: str = "system",
    author_label: str = "系统",
) -> CaseTaskComment:
    comment = CaseTaskComment(
        case_id=case_id,
        task_id=task_id,
        message=message.strip(),
        author_type=author_type,  # type: ignore[arg-type]
        author_label=author_label,
    )
    store.add("case_task_comments", comment)
    return comment


def validate_task_dependencies(case_id: str, task_id: str, depends_on_task_ids: list[str]) -> list[str]:
    normalized: list[str] = []
    for depends_on_id in depends_on_task_ids:
        if depends_on_id == task_id:
            raise HTTPException(status_code=409, detail="Task cannot depend on itself")
        dependency = store.get("case_tasks", depends_on_id, CaseTask)
        if not dependency or dependency.case_id != case_id:
            raise HTTPException(status_code=404, detail=f"Dependency task not found: {depends_on_id}")
        if depends_on_id not in normalized:
            normalized.append(depends_on_id)
    return normalized


def blocked_by_task_ids(case_id: str, task: CaseTask) -> list[str]:
    blocked_ids = []
    for depends_on_id in task.depends_on_task_ids:
        dependency = store.get("case_tasks", depends_on_id, CaseTask)
        if dependency and dependency.case_id == case_id and dependency.status != "done":
            blocked_ids.append(depends_on_id)
    return blocked_ids


def task_to_detail(task: CaseTask) -> dict[str, object]:
    payload = task.model_dump()
    payload["comments"] = [comment.model_dump() for comment in task_comments(task.case_id, task.id)]
    payload["research_results"] = [
        result.model_dump() for result in research_results(task.case_id, task.id)
    ]
    payload["blocked_by_task_ids"] = blocked_by_task_ids(task.case_id, task)
    return payload


def extract_research_keywords(case: Case, query: str, memories: list[CaseMemory]) -> list[str]:
    source = " ".join([case.title, case.summary, query, *[memory.content for memory in memories[:8]]])
    candidates = [
        "劳动合同",
        "拖欠工资",
        "被迫离职",
        "合同解除",
        "违约责任",
        "付款期限",
        "借款",
        "诉讼时效",
        "股权",
        "交通事故",
        "离婚",
        "抚养",
        "房屋买卖",
        "刑事拘留",
        "证据",
        "仲裁",
        "管辖",
        "违约金",
    ]
    keywords = [item for item in candidates if item in source]
    if case.case_type == "labor":
        keywords.extend(["劳动争议", "工资支付", "劳动仲裁"])
    elif case.case_type == "contract":
        keywords.extend(["合同纠纷", "履行期限", "违约责任"])
    elif case.case_type == "debt":
        keywords.extend(["民间借贷", "借款合意", "转账凭证"])
    elif case.case_type == "company":
        keywords.extend(["公司商事", "股东权利", "公司决议"])
    elif case.case_type == "criminal":
        keywords.extend(["刑事程序", "取保候审", "证据合法性"])
    else:
        keywords.append(case_type_label(case.case_type))
    seen: set[str] = set()
    return [item for item in keywords if not (item in seen or seen.add(item))][:8]


SEARCH_MARKUP_RE = re.compile(r"<[^>]+>")


def legal_search_client() -> DelilegalSearchClient:
    return DelilegalSearchClient()


def clean_search_text(value: Any) -> str:
    return " ".join(unescape(SEARCH_MARKUP_RE.sub("", str(value or ""))).split())


def truncate_search_text(value: Any, limit: int = 180) -> str:
    text = clean_search_text(value)
    if len(text) <= limit:
        return text
    return f"{text[:limit].rstrip()}..."


def metadata_keywords(value: Any) -> list[str]:
    if not isinstance(value, list):
        return []
    cleaned = [clean_search_text(item) for item in value]
    return [item for item in cleaned if item]


def normalize_highlights(value: Any) -> list[str]:
    if not isinstance(value, list):
        return []
    points: list[str] = []
    for item in value:
        if isinstance(item, str):
            text = truncate_search_text(item)
        elif isinstance(item, dict):
            text = truncate_search_text(item.get("content") or item.get("text") or item.get("fragment") or item)
        else:
            text = truncate_search_text(item)
        if text:
            points.append(text)
    return points


def ensure_search_api_success(response: dict[str, Any], label: str) -> None:
    if response.get("success"):
        return
    code = response.get("code")
    msg = response.get("msg") or "未知错误"
    raise LegalSearchApiError(f"{label}失败：{msg}", code=int(code) if isinstance(code, int) else None, raw_payload=response)


def research_relevance(index: int) -> float:
    return max(0.5, 0.95 - index * 0.03)


def build_similar_case_results(
    *,
    run: LegalResearchRun,
    case: Case,
    api_response: dict[str, Any],
) -> list[LegalResearchResult]:
    results: list[LegalResearchResult] = []
    for index, item in enumerate(api_response.get("data", [])):
        if not isinstance(item, dict):
            continue
        key_points = [
            point
            for point in [
                f"案由：{clean_search_text(item.get('cause'))}" if item.get("cause") else "",
                f"裁判类型：{clean_search_text(item.get('judgement_type'))}" if item.get("judgement_type") else "",
                f"审级：{clean_search_text(item.get('level_of_trial'))}" if item.get("level_of_trial") else "",
                f"裁判摘要：{truncate_search_text(item.get('content'))}" if item.get("content") else "",
            ]
            if point
        ][:5]
        external_id = clean_search_text(item.get("id")) or None
        reference_parts = [
            clean_search_text(item.get("case_number")),
            clean_search_text(item.get("judgement_date")),
            clean_search_text(item.get("publish_type_name")),
        ]
        results.append(
            LegalResearchResult(
                run_id=run.id,
                case_id=case.id,
                task_id=run.task_id,
                result_type="similar_case",
                external_id=external_id,
                title=clean_search_text(item.get("title")) or "未命名类案",
                source="法狗狗案例库",
                reference=" · ".join(part for part in reference_parts if part) or (external_id or ""),
                court_or_authority=clean_search_text(item.get("court")),
                relevance_score=research_relevance(index),
                key_points=key_points or ["真实案例库返回结果，建议律师复核裁判要旨、法院层级和事实相似度。"],
                metadata={
                    "provider": "delilegal",
                    "query_id": api_response.get("query_id"),
                    "api_code": api_response.get("code"),
                    "case_type": item.get("case_type"),
                    "cause": item.get("cause"),
                    "judgement_type": item.get("judgement_type"),
                    "judgement_date": item.get("judgement_date"),
                    "case_number": item.get("case_number"),
                    "level_of_trial": item.get("level_of_trial"),
                    "publish_type": item.get("publish_type"),
                    "publish_type_name": item.get("publish_type_name"),
                },
            )
        )
    return results


def build_regulation_results(
    *,
    run: LegalResearchRun,
    case: Case,
    api_response: dict[str, Any],
) -> list[LegalResearchResult]:
    results: list[LegalResearchResult] = []
    for index, item in enumerate(api_response.get("data", [])):
        if not isinstance(item, dict):
            continue
        highlights = normalize_highlights(item.get("highlights"))
        fallback_points = [
            point
            for point in [
                f"文号：{clean_search_text(item.get('issued_no'))}" if item.get("issued_no") else "",
                f"层级：{clean_search_text(item.get('level_name'))}" if item.get("level_name") else "",
                f"时效性：{clean_search_text(item.get('timeliness_name'))}" if item.get("timeliness_name") else "",
                f"发布日期：{clean_search_text(item.get('publish_date'))}" if item.get("publish_date") else "",
            ]
            if point
        ]
        external_id = clean_search_text(item.get("id")) or None
        reference_parts = [
            clean_search_text(item.get("issued_no")),
            clean_search_text(item.get("publish_date")),
            clean_search_text(item.get("timeliness_name")),
        ]
        results.append(
            LegalResearchResult(
                run_id=run.id,
                case_id=case.id,
                task_id=run.task_id,
                result_type="regulation",
                external_id=external_id,
                title=clean_search_text(item.get("title")) or "未命名法规",
                source="法狗狗法规库",
                reference=" · ".join(part for part in reference_parts if part) or (external_id or ""),
                court_or_authority=clean_search_text(item.get("publisher_name")),
                relevance_score=research_relevance(index),
                key_points=(highlights or fallback_points)[:5]
                or ["真实法规库返回结果，建议律师点击法规详情核验全文、时效和适用范围。"],
                metadata={
                    "provider": "delilegal",
                    "query_id": api_response.get("query_id"),
                    "api_code": api_response.get("code"),
                    "issued_no": item.get("issued_no"),
                    "publish_date": item.get("publish_date"),
                    "publisher_name": item.get("publisher_name"),
                    "active_date": item.get("active_date"),
                    "timeliness_name": item.get("timeliness_name"),
                    "level_name": item.get("level_name"),
                    "highlights": item.get("highlights", []),
                },
            )
        )
    return results


def mark_research_task_failed(task: CaseTask, run: LegalResearchRun, failure_reason: str) -> CaseTask:
    run.status = "failed"
    run.failure_reason = failure_reason
    run.summary = failure_reason
    run.completed_at = now_iso()
    store.update("legal_research_runs", run)
    task.status = "blocked"
    task.result_summary = f"真实 API 调用失败：{failure_reason}"
    task.metadata = {
        **task.metadata,
        "research_run_id": run.id,
        "api_provider": "delilegal",
        "api_success": False,
        "api_failure_reason": failure_reason,
    }
    task.updated_at = now_iso()
    store.update("case_tasks", task)
    add_task_comment(case_id=task.case_id, task_id=task.id, message=task.result_summary)
    return task


def execute_research_task(case: Case, task: CaseTask, payload: TaskExecuteRequest) -> CaseTask:
    query = clean_search_text(payload.query or task.metadata.get("query") or task.description or case.summary or case.title)
    search_type = "similar_case" if task.task_type == "similar_case_search" else "regulation"
    explicit_keywords = metadata_keywords(payload.keywords) or metadata_keywords(task.metadata.get("keywords"))
    keywords = explicit_keywords or extract_research_keywords(case, query, case_memories(case.id))
    search_text = query or " ".join(keywords) or case.title
    run = LegalResearchRun(
        case_id=case.id,
        task_id=task.id,
        search_type=search_type,  # type: ignore[arg-type]
        query=search_text,
        keywords=keywords,
    )
    store.add("legal_research_runs", run)

    try:
        client = legal_search_client()
        if search_type == "similar_case":
            api_response = client.search_cases(
                keyword=search_text,
                page_no=payload.page_no,
                page_size=payload.page_size,
                sort_field=payload.sort_field,
                sort_order=payload.sort_order,
            )
            ensure_search_api_success(api_response, "类案检索")
            results = build_similar_case_results(run=run, case=case, api_response=api_response)
        else:
            law_keywords = explicit_keywords or [search_text]
            api_response = client.search_laws(
                keywords=law_keywords,
                field_name=payload.field_name,
                page_no=payload.page_no,
                page_size=payload.page_size,
                sort_field=payload.sort_field,
                sort_order=payload.sort_order,
            )
            ensure_search_api_success(api_response, "法规检索")
            results = build_regulation_results(run=run, case=case, api_response=api_response)
    except LegalSearchApiError as exc:
        return mark_research_task_failed(task, run, str(exc))
    except Exception as exc:
        return mark_research_task_failed(task, run, f"检索服务异常：{exc}")

    for result in results:
        store.add("legal_research_results", result)
    run.status = "completed"
    run.result_count = len(results)
    run.summary = (
        f"已调用法狗狗真实 API，返回 {len(results)} 条"
        f"{'类案' if search_type == 'similar_case' else '法规'}结果"
        f"（总数 {api_response.get('total_count', 0)}，queryId {api_response.get('query_id') or '无'}）。"
    )
    run.completed_at = now_iso()
    store.update("legal_research_runs", run)
    task.status = "waiting_owner_review"
    task.result_summary = run.summary
    task.metadata = {
        **task.metadata,
        "research_run_id": run.id,
        "keywords": keywords,
        "query": search_text,
        "api_provider": "delilegal",
        "api_success": True,
        "api_query_id": api_response.get("query_id"),
        "api_total_count": api_response.get("total_count"),
        "api_total_page": api_response.get("total_page"),
        "api_code": api_response.get("code"),
        "api_msg": api_response.get("msg"),
        "page_no": payload.page_no,
        "page_size": payload.page_size,
        "sort_field": payload.sort_field,
        "sort_order": payload.sort_order,
        "field_name": payload.field_name,
    }
    task.updated_at = now_iso()
    store.update("case_tasks", task)
    store.add(
        "case_memories",
        CaseMemory(
            case_id=case.id,
            kind="note",
            content=f"{task.title}：{run.summary}",
            confidence=0.82,
            source_ref=run.id,
        ),
    )
    add_task_comment(case_id=case.id, task_id=task.id, message=run.summary)
    return task


def build_document_draft_content(case: Case, task: CaseTask, payload: TaskExecuteRequest) -> str:
    if payload.content_text and payload.content_text.strip():
        return payload.content_text.strip()
    memories = case_memories(case.id)
    facts = [memory.content for memory in memories if memory.kind in {"fact", "timeline", "evidence"}]
    uncertainties = [memory.content for memory in memories if memory.kind == "uncertainty"]
    research = research_results(case.id)
    research_lines = []
    for result in research[:4]:
        points = "；".join(result.key_points[:2])
        research_lines.append(f"- {result.title}：{points}")
    facts_text = "\n".join(f"- {item}" for item in facts[:8]) or "- 暂无充分事实记录。"
    uncertainty_text = "\n".join(f"- {item}" for item in uncertainties[:5]) or "- 暂未登记重大不确定点。"
    research_text = "\n".join(research_lines) or "- 暂无检索结果，建议先执行类案或法规检索任务。"
    return f"""文书草稿：{payload.title or task.title}
案件名称：{case.title}
案件类型：{case_type_label(case.case_type)}

一、事实基础
{facts_text}

二、法律与类案检索参考
{research_text}

三、初步处理意见
{case_rule_hint(case.case_type)}
结合现有材料，建议先确认关键事实、证据完整性和客户目标，再确定正式对外文本。

四、风险与待补充事项
{uncertainty_text}

五、人工复核
本草稿由文档撰写任务生成，应由律师复核事实来源、法律依据、证据引用和表述边界后再提交或发送。
"""


def execute_document_review_task(case: Case, task: CaseTask, payload: TaskExecuteRequest) -> CaseTask:
    document_id = payload.document_id or task.document_id
    if not document_id:
        raise HTTPException(status_code=422, detail="document_id is required for document review")
    document = store.get("legal_documents", document_id, LegalDocument)
    if not document or document.case_id not in {None, case.id}:
        raise HTTPException(status_code=404, detail="Document not found")
    revisions = sorted(
        store.filter("legal_document_revisions", LegalDocumentRevision, document_id=document_id),
        key=lambda item: item.version_number,
    )
    if len(revisions) < 2 and not (payload.base_revision_id or task.base_revision_id):
        raise HTTPException(status_code=422, detail="At least two revisions are required for document review")
    base_id = payload.base_revision_id or task.base_revision_id or revisions[-2].id
    target_id = payload.target_revision_id or task.target_revision_id or revisions[-1].id
    base = store.get("legal_document_revisions", base_id, LegalDocumentRevision)
    target = store.get("legal_document_revisions", target_id, LegalDocumentRevision)
    if not base or base.document_id != document_id:
        raise HTTPException(status_code=404, detail="Base revision not found")
    if not target or target.document_id != document_id:
        raise HTTPException(status_code=404, detail="Target revision not found")
    diff = LegalDocumentDiff(
        document_id=document.id,
        base_revision_id=base.id,
        target_revision_id=target.id,
        segments=build_char_diff(base.content_text, target.content_text),
        paragraph_changes=build_paragraph_diff(base.content_text, target.content_text),
        risk_summary=summarize_legal_risks(base.content_text, target.content_text),
    )
    existing_diff = store.filter(
        "legal_document_diffs",
        LegalDocumentDiff,
        document_id=document.id,
        base_revision_id=base.id,
        target_revision_id=target.id,
    )
    if existing_diff:
        diff = existing_diff[-1]
    else:
        store.add("legal_document_diffs", diff)
    existing_analysis = store.filter(
        "legal_document_analyses",
        LegalDocumentAnalysis,
        document_id=document.id,
        base_revision_id=base.id,
        target_revision_id=target.id,
    )
    analysis = existing_analysis[-1] if existing_analysis else LegalDocumentAnalysis(
        document_id=document.id,
        base_revision_id=base.id,
        target_revision_id=target.id,
        source="rule_fallback",
        risk_level="medium" if diff.risk_summary else "low",
        ambiguities=["发现文本变化，请人工核对是否影响权利义务、履行期限或争议解决。"] if diff.risk_summary else [],
        risk_points=diff.risk_summary,
        suggestions=["建议逐条核对版本差异，并确认是否需要回退、接受或另开修订分支。"],
        manual_review_checklist=[
            "核对变更是否影响金额、期限、责任、解除、管辖或争议解决。",
            "核对删除内容是否弱化我方权利或扩大对方免责。",
            "核对新增内容是否需要客户确认或主任律师复核。",
        ],
    )
    if not existing_analysis:
        store.add("legal_document_analyses", analysis)
    task.status = "waiting_owner_review"
    task.document_id = document.id
    task.base_revision_id = base.id
    task.target_revision_id = target.id
    task.result_summary = f"已审查 {document.title} v{base.version_number} -> v{target.version_number}，风险等级：{analysis.risk_level}，风险点 {len(analysis.risk_points)} 项。"
    task.metadata = {**task.metadata, "diff_id": diff.id, "analysis_id": analysis.id}
    task.updated_at = now_iso()
    store.update("case_tasks", task)
    add_task_comment(case_id=case.id, task_id=task.id, message=task.result_summary)
    return task


def execute_document_drafting_task(case: Case, task: CaseTask, payload: TaskExecuteRequest) -> CaseTask:
    content = build_document_draft_content(case, task, payload)
    document_id = payload.document_id or task.document_id
    title = str(payload.title or task.metadata.get("title") or f"{task.title} 草稿")
    change_summary = payload.change_summary or f"文档撰写任务 {task.id} 生成草稿"
    draft_branch_name = "main"
    if document_id:
        document = store.get("legal_documents", document_id, LegalDocument)
        if not document or document.case_id not in {None, case.id}:
            raise HTTPException(status_code=404, detail="Document not found")
        revisions = sorted(
            store.filter("legal_document_revisions", LegalDocumentRevision, document_id=document.id),
            key=lambda item: item.version_number,
        )
        base_revision_id = payload.base_revision_id or task.base_revision_id or document.current_revision_id or (revisions[-1].id if revisions else None)
        if not base_revision_id:
            raise HTTPException(status_code=422, detail="Base revision is required")
        branch_name = f"draft/{task.id.replace('task_', '')[:8]}"
        draft_branch_name = branch_name
        branches = document_branches(document.id)
        branch = next((item for item in branches if item.name == branch_name), None)
        if not branch:
            base_revision = store.get("legal_document_revisions", base_revision_id, LegalDocumentRevision)
            if not base_revision or base_revision.document_id != document.id:
                raise HTTPException(status_code=404, detail="Base revision not found")
            branch = LegalDocumentBranch(
                document_id=document.id,
                name=branch_name,
                head_revision_id=base_revision.id,
                base_revision_id=base_revision.id,
            )
            store.add("legal_document_branches", branch)
        revision = create_revision_on_branch(
            document=document,
            branch=branch,
            content_text=content,
            source_filename=f"{title}.docx",
            author_type="agent",
            change_summary=change_summary,
        )
    else:
        document = create_text_document(
            case_id=case.id,
            title=title,
            document_type="pleading",
            content_text=content,
            source_filename=f"{title}.docx",
            change_summary=change_summary,
            author_type="agent",
        )
        revision = latest_document_revision(document.id)
        if not revision:
            raise HTTPException(status_code=500, detail="Draft revision was not created")
    task.status = "waiting_owner_review"
    task.output_document_id = document.id
    task.output_revision_id = revision.id
    task.document_id = document.id
    task.result_summary = f"已生成文书草稿：{document.title}，版本 v{revision.version_number}。"
    task.metadata = {**task.metadata, "draft_branch": draft_branch_name}
    task.updated_at = now_iso()
    store.update("case_tasks", task)
    add_task_comment(case_id=case.id, task_id=task.id, message=task.result_summary)
    return task


def build_agent_architecture() -> AgentArchitecture:
    agents = store.list("legal_agents", LegalAgent)
    by_role = {agent.role: agent for agent in agents}
    dispatcher = by_role.get("dispatch_agent") or by_role.get("managing_lawyer")
    if not dispatcher:
        raise HTTPException(status_code=500, detail="Agent architecture is not initialized")
    return AgentArchitecture(
        dispatcher=dispatcher,
        groups=[
            AgentGroup(
                id="core_business",
                title="核心业务 Agent",
                description="承接专业法律判断、深度推理和文书交付。",
                agent_roles=[
                    "core_business_agent",
                    "managing_lawyer",
                    "handling_lawyer",
                    "litigation_strategist",
                    "legal_researcher",
                    "drafting_lawyer",
                ],
                departments=CORE_DEPARTMENTS,
            ),
            AgentGroup(
                id="client_service",
                title="客户服务 Agent",
                description="负责法律咨询、接案报价、投诉处理和客户沟通节奏。",
                agent_roles=["client_service_agent", "reception_lawyer"],
                departments=CLIENT_SERVICE_DEPARTMENTS,
            ),
            AgentGroup(
                id="compliance_review",
                title="合规审查 Agent",
                description="负责收案审批、利益冲突审查、风险评估和人工复核点。",
                agent_roles=["compliance_review_agent", "quality_control", "contract_reviewer"],
                departments=COMPLIANCE_REVIEW_DEPARTMENTS,
            ),
            AgentGroup(
                id="archive_management",
                title="档案管理 Agent",
                description="负责案卷归档、档案查询、证据目录、文件版本和过程留痕。",
                agent_roles=["archive_management_agent", "case_secretary"],
                departments=ARCHIVE_MANAGEMENT_DEPARTMENTS,
            ),
        ],
    )


@app.get("/api/health")
async def health() -> dict[str, object]:
    return {"ok": True, "service": "lvzhijie-backend"}


@app.get("/api/dashboard/summary")
async def dashboard_summary() -> dict[str, object]:
    cases = store.list("cases", Case)
    conversations = store.list("wechat_conversations", WechatConversation)
    messages = store.list("wechat_messages", WechatMessage)
    documents = store.list("legal_documents", LegalDocument)
    runs = store.list("legal_reasoning_runs", LegalReasoningRun)
    reply_jobs = store.list("reply_jobs", LegalReplyJob)
    legal_research_runs = store.list("legal_research_runs", LegalResearchRun)
    legal_research_results = store.list("legal_research_results", LegalResearchResult)
    return {
        "cases": len(cases),
        "open_cases": len([case for case in cases if case.status != "closed"]),
        "conversations": len(conversations),
        "unread": sum(item.unread_count for item in conversations),
        "messages": len(messages),
        "documents": len(documents),
        "reasoning_runs": len(runs),
        "research_runs": len(legal_research_runs),
        "research_results": len(legal_research_results),
        "reply_jobs": len(reply_jobs),
        "queued_reply_jobs": len([job for job in reply_jobs if job.status in {"queued", "reasoning"}]),
        "openclaw": (await OpenClawWechatAdapter(connection()).get_status()).model_dump(),
    }


@app.get("/api/activity")
async def list_activity() -> list[ActivityEvent]:
    return sorted(
        store.list("activity_events", ActivityEvent),
        key=lambda item: item.created_at,
        reverse=True,
    )


@app.get("/api/openclaw/connection")
async def get_openclaw_connection() -> OpenClawConnection:
    return connection()


@app.put("/api/openclaw/connection")
async def update_openclaw_connection(payload: OpenClawConnection) -> OpenClawConnection:
    payload.last_checked_at = now_iso()
    store.upsert_single("openclaw_connection", payload)
    record("openclaw.connection.updated", "更新 OpenClaw 微信通道配置")
    return payload


@app.get("/api/openclaw/status")
async def openclaw_status() -> dict[str, object]:
    status = await OpenClawWechatAdapter(connection()).get_status()
    updated = connection()
    updated.last_checked_at = status.checked_at
    store.upsert_single("openclaw_connection", updated)
    return status.model_dump()


@app.post("/api/openclaw/sync")
async def sync_openclaw_messages() -> dict[str, object]:
    if using_mock_wechat():
        mock_wechat.sync_to_json_store(store)
        conversations = mock_wechat.list_conversations()
        total_messages = sum(
            len(mock_wechat.list_messages(c["id"])) for c in conversations
        )
        return {
            "ok": True,
            "sessions": len(conversations),
            "messages": total_messages,
            "errors": [],
            "last_sync_at": now_iso(),
        }

    adapter = OpenClawWechatAdapter(connection())
    synced_sessions = 0
    synced_messages = 0
    errors: list[str] = []
    try:
        sessions = await adapter.list_sessions()
    except Exception as exc:  # noqa: BLE001 - sync should report, not crash the workspace
        errors.append(str(exc))
        sessions = []

    contacts = store.list("wechat_contacts", WechatContact)
    conversations = store.list("wechat_conversations", WechatConversation)
    messages = store.list("wechat_messages", WechatMessage)
    contacts_by_openclaw = {item.openclaw_contact_id: item for item in contacts}
    conversations_by_openclaw = {item.openclaw_conversation_id: item for item in conversations}
    existing_message_keys = {
        _message_dedupe_key(message)
        for message in messages
    }

    for raw_session in sessions:
        session_key = adapter.session_key(raw_session)
        label = adapter.session_label(raw_session)
        contact = contacts_by_openclaw.get(session_key)
        if contact is None:
            contact = WechatContact(
                openclaw_contact_id=session_key,
                display_name=label,
                remark="OpenClaw 微信会话",
                last_seen_at=now_iso(),
            )
            store.add("wechat_contacts", contact)
            contacts_by_openclaw[session_key] = contact
        else:
            contact.display_name = label or contact.display_name
            contact.last_seen_at = now_iso()
            store.update("wechat_contacts", contact)

        conversation = conversations_by_openclaw.get(session_key)
        if conversation is None:
            conversation = WechatConversation(
                openclaw_conversation_id=session_key,
                contact_id=contact.id,
                status="open",
                auto_reply_source="openclaw",
            )
            store.add("wechat_conversations", conversation)
            conversations_by_openclaw[session_key] = conversation
        synced_sessions += 1

        try:
            history = await adapter.get_session_history(session_key)
        except Exception as exc:  # noqa: BLE001
            errors.append(f"{session_key}: {exc}")
            continue
        for raw_message in history:
            normalized = adapter.normalize_history_message(
                raw_message,
                conversation_id=conversation.id,
            )
            key = _message_dedupe_key(normalized)
            if key in existing_message_keys or not normalized.content.strip():
                continue
            store.add("wechat_messages", normalized)
            existing_message_keys.add(key)
            synced_messages += 1
            conversation.last_message_at = normalized.created_at
        store.update("wechat_conversations", conversation)

    updated = connection()
    updated.last_sync_at = now_iso()
    store.upsert_single("openclaw_connection", updated)
    record(
        "wechat.sync",
        "同步 OpenClaw 微信聊天记录",
        f"会话 {synced_sessions} 个，消息 {synced_messages} 条。",
    )
    return {
        "ok": not errors,
        "sessions": synced_sessions,
        "messages": synced_messages,
        "errors": errors,
        "last_sync_at": updated.last_sync_at,
    }


@app.get("/api/openclaw/sessions")
async def list_openclaw_sessions() -> dict[str, object]:
    adapter = OpenClawWechatAdapter(connection())
    try:
        sessions = await adapter.list_sessions()
        return {"ok": True, "sessions": sessions}
    except Exception as exc:  # noqa: BLE001
        return {"ok": False, "sessions": [], "error": str(exc)}


@app.get("/api/wechat/conversations")
async def list_wechat_conversations() -> list[dict[str, object]]:
    sync_mock_wechat_if_needed()
    contacts = {item.id: item for item in store.list("wechat_contacts", WechatContact)}
    conversations = store.list("wechat_conversations", WechatConversation)
    return [
        {
            **conversation.model_dump(),
            "contact": contacts.get(conversation.contact_id).model_dump()
            if contacts.get(conversation.contact_id)
            else None,
        }
        for conversation in sorted(
            conversations,
            key=lambda item: item.last_message_at or "",
            reverse=True,
        )
    ]


@app.get("/api/wechat/conversations/{conversation_id}/messages")
async def get_wechat_messages(conversation_id: str) -> list[WechatMessage]:
    sync_mock_wechat_if_needed()
    return sorted(
        store.filter("wechat_messages", WechatMessage, conversation_id=conversation_id),
        key=lambda item: item.created_at,
    )


@app.delete("/api/wechat/conversations/{conversation_id}")
async def delete_wechat_conversation(conversation_id: str) -> dict[str, object]:
    sync_mock_wechat_if_needed()
    conversation = store.get("wechat_conversations", conversation_id, WechatConversation)
    if not conversation:
        raise HTTPException(status_code=404, detail="Conversation not found")
    if using_mock_wechat():
        mock_wechat.delete_conversation(conversation_id)
    result = delete_wechat_conversation_rows(conversation)
    record(
        "wechat.conversation.deleted",
        "删除微信会话",
        conversation.openclaw_conversation_id,
        entity_type="conversation",
        entity_id=conversation_id,
    )
    return {"ok": True, **result}


@app.post("/api/wechat/conversations/{conversation_id}/send")
async def send_wechat_message(conversation_id: str, payload: SendMessageRequest) -> WechatMessage:
    conversation = store.get("wechat_conversations", conversation_id, WechatConversation)
    if not conversation:
        raise HTTPException(status_code=404, detail="Conversation not found")
    if using_mock_wechat():
        new_msg = mock_wechat.append_message(
            conversation_id=conversation_id,
            sender="owner",
            content=payload.content,
            source="manual",
        )
        message = WechatMessage.model_validate(new_msg)
        mock_wechat.sync_to_json_store(store)
        record(
            "wechat.sent_mock",
            "通过演示模式发送微信消息",
            payload.content[:120],
            entity_type="conversation",
            entity_id=conversation_id,
        )
        return message
    message = await OpenClawWechatAdapter(connection()).send_wechat_message(
        conversation_id=conversation.openclaw_conversation_id,
        content=payload.content,
    )
    message.conversation_id = conversation_id
    conversation.last_message_at = message.created_at
    store.add("wechat_messages", message)
    store.update("wechat_conversations", conversation)
    record(
        "wechat.sent_via_openclaw",
        "通过 OpenClaw 发送微信消息",
        payload.content[:120],
        entity_type="conversation",
        entity_id=conversation_id,
    )
    return message


def _message_dedupe_key(message: WechatMessage) -> str:
    if message.openclaw_message_id:
        return f"oc:{message.conversation_id}:{message.openclaw_message_id}"
    return (
        f"fp:{message.conversation_id}:{message.sender}:"
        f"{message.created_at}:{message.content[:160]}"
    )


@app.post("/api/wechat/conversations/{conversation_id}/case")
async def create_case_from_conversation(
    conversation_id: str,
    payload: CreateCaseFromConversationRequest,
) -> Case:
    sync_mock_wechat_if_needed()
    conversation = store.get("wechat_conversations", conversation_id, WechatConversation)
    if not conversation:
        raise HTTPException(status_code=404, detail="Conversation not found")
    if conversation.case_id:
        existing_case = store.get("cases", conversation.case_id, Case)
        if existing_case:
            return existing_case
    contact = store.get("wechat_contacts", conversation.contact_id, WechatContact)
    messages = sorted(
        store.filter("wechat_messages", WechatMessage, conversation_id=conversation_id),
        key=lambda item: item.created_at,
    )
    analysis = await analyze_conversation_for_case(
        contact=contact,
        messages=messages,
        requested_title=payload.title,
        requested_case_type=payload.case_type if payload.case_type != "other" else None,
    )
    case = Case(
        title=str(analysis["title"]),
        case_type=analysis["case_type"],  # type: ignore[arg-type]
        status="collecting_info",
        summary=str(analysis["summary"]),
        wechat_contact_ref=conversation.contact_id,
        conversation_ref=conversation.id,
    )
    conversation.case_id = case.id
    store.add("cases", case)
    store.update("wechat_conversations", conversation)
    if using_mock_wechat():
        mock_wechat.update_conversation(conversation.id, {"case_id": case.id})
    for content in analysis.get("facts", []):
        store.add(
            "case_memories",
            CaseMemory(case_id=case.id, kind="fact", content=str(content), source_ref=conversation.id),
        )
    for content in analysis.get("evidence", []):
        store.add(
            "case_memories",
            CaseMemory(case_id=case.id, kind="evidence", content=str(content), source_ref=conversation.id),
        )
    for content in analysis.get("uncertainties", []):
        store.add(
            "case_memories",
            CaseMemory(case_id=case.id, kind="uncertainty", content=str(content), source_ref=conversation.id),
        )
    for title in analysis.get("suggested_tasks", [])[:5]:
        store.add(
            "case_tasks",
            CaseTask(
                case_id=case.id,
                title=str(title),
                assigned_agent_role="案件秘书 Agent",
            ),
        )
    record("case.created", "从微信会话创建案件", case.title, entity_type="case", entity_id=case.id)
    return case


@app.post("/api/wechat/conversations/{conversation_id}/bind-case")
async def bind_conversation_to_case(
    conversation_id: str,
    payload: BindConversationCaseRequest,
) -> Case:
    sync_mock_wechat_if_needed()
    conversation = store.get("wechat_conversations", conversation_id, WechatConversation)
    if not conversation:
        raise HTTPException(status_code=404, detail="Conversation not found")
    case = store.get("cases", payload.case_id, Case)
    if not case:
        raise HTTPException(status_code=404, detail="Case not found")
    conversation.case_id = case.id
    case.conversation_ref = conversation.id
    case.wechat_contact_ref = conversation.contact_id
    case.updated_at = now_iso()
    store.update("wechat_conversations", conversation)
    store.update("cases", case)
    if using_mock_wechat():
        mock_wechat.update_conversation(conversation.id, {"case_id": case.id})
    record("case.bound_wechat", "绑定微信会话到案件", case.title, entity_type="case", entity_id=case.id)
    return case


@app.get("/api/cases")
async def list_cases() -> list[Case]:
    return sorted(store.list("cases", Case), key=lambda item: item.updated_at, reverse=True)


@app.post("/api/cases")
async def create_case(payload: CaseCreateRequest) -> Case:
    case = Case(**payload.model_dump())
    store.add("cases", case)
    record("case.created", "创建案件", case.title, entity_type="case", entity_id=case.id)
    return case


@app.get("/api/cases/{case_id}")
async def get_case(case_id: str) -> dict[str, object]:
    case = store.get("cases", case_id, Case)
    if not case:
        raise HTTPException(status_code=404, detail="Case not found")
    return {
        "case": case,
        "tasks": [task_to_detail(task) for task in case_tasks(case_id)],
        "task_comments": task_comments(case_id),
        "research_runs": research_runs(case_id),
        "research_results": research_results(case_id),
        "memories": store.filter("case_memories", CaseMemory, case_id=case_id),
        "documents": store.filter("legal_documents", LegalDocument, case_id=case_id),
        "reasoning_runs": store.filter("legal_reasoning_runs", LegalReasoningRun, case_id=case_id),
        "follow_up_questions": store.filter("follow_up_questions", FollowUpQuestion, case_id=case_id),
        "reply_jobs": store.filter("reply_jobs", LegalReplyJob, case_id=case_id),
        "messages": store.filter("wechat_messages", WechatMessage, conversation_id=case.conversation_ref)
        if case.conversation_ref
        else [],
    }


@app.delete("/api/cases/{case_id}")
async def delete_case(case_id: str) -> dict[str, object]:
    case = store.get("cases", case_id, Case)
    if not case:
        raise HTTPException(status_code=404, detail="Case not found")

    for document in store.filter("legal_documents", LegalDocument, case_id=case_id):
        delete_document_rows(document.id)
    store.remove_where("case_tasks", lambda row: row.get("case_id") == case_id)
    store.remove_where("case_task_comments", lambda row: row.get("case_id") == case_id)
    store.remove_where("legal_research_runs", lambda row: row.get("case_id") == case_id)
    store.remove_where("legal_research_results", lambda row: row.get("case_id") == case_id)
    store.remove_where("case_memories", lambda row: row.get("case_id") == case_id)
    store.remove_where("legal_reasoning_runs", lambda row: row.get("case_id") == case_id)
    store.remove_where("follow_up_questions", lambda row: row.get("case_id") == case_id)
    store.remove_where("reply_jobs", lambda row: row.get("case_id") == case_id)
    store.remove_where("activity_events", lambda row: row.get("entity_type") == "case" and row.get("entity_id") == case_id)

    for conversation in store.list("wechat_conversations", WechatConversation):
        if conversation.case_id == case_id:
            conversation.case_id = None
            store.update("wechat_conversations", conversation)

    store.delete("cases", case_id)
    record("case.deleted", "删除案件", case.title)
    return {"ok": True}


@app.post("/api/cases/{case_id}/tasks")
async def create_task(case_id: str, payload: TaskCreateRequest) -> CaseTask:
    ensure_case_exists(case_id)
    task = CaseTask(case_id=case_id, **payload.model_dump())
    task.depends_on_task_ids = validate_task_dependencies(case_id, task.id, task.depends_on_task_ids)
    store.add("case_tasks", task)
    if task.description.strip():
        add_task_comment(
            case_id=case_id,
            task_id=task.id,
            message=f"任务说明：{task.description.strip()}",
            author_type="system",
            author_label="任务中心",
        )
    record("case.task.created", "创建案件任务", task.title, entity_type="case", entity_id=case_id)
    return task


@app.patch("/api/cases/{case_id}/tasks/{task_id}")
async def update_task(case_id: str, task_id: str, payload: TaskUpdateRequest) -> CaseTask:
    ensure_case_exists(case_id)
    task = get_case_task(case_id, task_id)
    updates = payload.model_dump(exclude_unset=True)
    comment = updates.pop("comment", None)
    if "depends_on_task_ids" in updates and updates["depends_on_task_ids"] is not None:
        updates["depends_on_task_ids"] = validate_task_dependencies(
            case_id,
            task_id,
            updates["depends_on_task_ids"],
        )
    target_status = updates.get("status")
    if target_status and target_status not in {"todo", "blocked", task.status, "done"}:
        blocking = blocked_by_task_ids(case_id, task)
        if blocking:
            raise HTTPException(
                status_code=409,
                detail={
                    "message": "Task is blocked by incomplete dependencies.",
                    "blocked_by_task_ids": blocking,
                },
            )
    if target_status == "done":
        blocking = blocked_by_task_ids(case_id, task)
        if blocking:
            raise HTTPException(
                status_code=409,
                detail={
                    "message": "Task cannot be completed before dependencies are done.",
                    "blocked_by_task_ids": blocking,
                },
            )
    for field, value in updates.items():
        setattr(task, field, value)
    task.updated_at = now_iso()
    store.update("case_tasks", task)
    if comment and str(comment).strip():
        add_task_comment(case_id=case_id, task_id=task_id, message=str(comment), author_type="owner", author_label="人工")
    record("case.task.updated", "更新案件任务", task.title, entity_type="case", entity_id=case_id)
    return task


@app.get("/api/cases/{case_id}/tasks/{task_id}/comments")
async def list_task_comments(case_id: str, task_id: str) -> list[CaseTaskComment]:
    ensure_case_exists(case_id)
    get_case_task(case_id, task_id)
    return task_comments(case_id, task_id)


@app.post("/api/cases/{case_id}/tasks/{task_id}/comments")
async def create_task_comment(
    case_id: str,
    task_id: str,
    payload: TaskCommentCreateRequest,
) -> CaseTaskComment:
    ensure_case_exists(case_id)
    get_case_task(case_id, task_id)
    if not payload.message.strip():
        raise HTTPException(status_code=422, detail="Comment cannot be empty")
    comment = add_task_comment(
        case_id=case_id,
        task_id=task_id,
        message=payload.message,
        author_type=payload.author_type,
        author_label=payload.author_label,
    )
    record("case.task.comment", "记录任务评论", payload.message[:120], entity_type="case", entity_id=case_id)
    return comment


@app.post("/api/cases/{case_id}/tasks/{task_id}/execute")
async def execute_task(case_id: str, task_id: str, payload: TaskExecuteRequest | None = None) -> CaseTask:
    case = ensure_case_exists(case_id)
    task = get_case_task(case_id, task_id)
    payload = payload or TaskExecuteRequest()
    blocking = blocked_by_task_ids(case_id, task)
    if blocking:
        raise HTTPException(
            status_code=409,
            detail={
                "message": "Task is blocked by incomplete dependencies.",
                "blocked_by_task_ids": blocking,
            },
        )
    task.status = "in_progress"
    task.updated_at = now_iso()
    store.update("case_tasks", task)
    if task.task_type in {"similar_case_search", "regulation_search"}:
        executed = execute_research_task(case, task, payload)
    elif task.task_type == "document_review":
        executed = execute_document_review_task(case, task, payload)
    elif task.task_type == "document_drafting":
        executed = execute_document_drafting_task(case, task, payload)
    else:
        task.status = "done"
        task.result_summary = task.result_summary or "通用任务已标记完成。"
        task.updated_at = now_iso()
        store.update("case_tasks", task)
        add_task_comment(case_id=case_id, task_id=task_id, message=task.result_summary)
        executed = task
    record("case.task.executed", "执行案件任务", executed.title, entity_type="case", entity_id=case_id)
    return executed


def raise_legal_search_http_error(exc: LegalSearchApiError) -> None:
    raise HTTPException(
        status_code=502,
        detail={
            "message": str(exc),
            "code": exc.code,
        },
    )


@app.post("/api/legal-research/cases/search")
async def direct_similar_case_search(payload: TaskExecuteRequest) -> dict[str, Any]:
    keyword = clean_search_text(payload.query) or " ".join(metadata_keywords(payload.keywords))
    if not keyword:
        raise HTTPException(status_code=422, detail="Search query cannot be empty")
    try:
        response = legal_search_client().search_cases(
            keyword=keyword,
            page_no=payload.page_no,
            page_size=payload.page_size,
            sort_field=payload.sort_field,
            sort_order=payload.sort_order,
        )
        ensure_search_api_success(response, "类案检索")
        return response
    except LegalSearchApiError as exc:
        raise_legal_search_http_error(exc)


@app.post("/api/legal-research/laws/search")
async def direct_law_search(payload: TaskExecuteRequest) -> dict[str, Any]:
    keywords = metadata_keywords(payload.keywords) or [clean_search_text(payload.query)]
    keywords = [item for item in keywords if item]
    if not keywords:
        raise HTTPException(status_code=422, detail="Search query cannot be empty")
    try:
        response = legal_search_client().search_laws(
            keywords=keywords,
            field_name=payload.field_name,
            page_no=payload.page_no,
            page_size=payload.page_size,
            sort_field=payload.sort_field,
            sort_order=payload.sort_order,
        )
        ensure_search_api_success(response, "法规检索")
        return response
    except LegalSearchApiError as exc:
        raise_legal_search_http_error(exc)


@app.get("/api/legal-research/laws/{law_id}/detail")
async def direct_law_detail(law_id: str, merge: bool = True) -> dict[str, Any]:
    try:
        response = legal_search_client().get_law_detail(law_id=law_id, merge=merge)
        ensure_search_api_success(response, "法规详情")
        return response
    except LegalSearchApiError as exc:
        raise_legal_search_http_error(exc)


@app.delete("/api/cases/{case_id}/tasks/{task_id}")
async def delete_task(case_id: str, task_id: str) -> dict[str, object]:
    task = get_case_task(case_id, task_id)
    store.delete("case_tasks", task_id)
    store.remove_where("case_task_comments", lambda row: row.get("task_id") == task_id)
    store.remove_where("legal_research_runs", lambda row: row.get("task_id") == task_id)
    store.remove_where("legal_research_results", lambda row: row.get("task_id") == task_id)
    for other in store.filter("case_tasks", CaseTask, case_id=case_id):
        if task_id in other.depends_on_task_ids:
            other.depends_on_task_ids = [item for item in other.depends_on_task_ids if item != task_id]
            other.updated_at = now_iso()
            store.update("case_tasks", other)
    record("case.task.deleted", "删除案件任务", task.title, entity_type="case", entity_id=case_id)
    return {"ok": True}


@app.post("/api/cases/{case_id}/memories")
async def create_memory(case_id: str, payload: MemoryCreateRequest) -> CaseMemory:
    if not store.get("cases", case_id, Case):
        raise HTTPException(status_code=404, detail="Case not found")
    memory = CaseMemory(case_id=case_id, **payload.model_dump())
    store.add("case_memories", memory)
    record("case.memory.created", "写入案件记忆", memory.content[:120], entity_type="case", entity_id=case_id)
    return memory


@app.delete("/api/cases/{case_id}/memories/{memory_id}")
async def delete_memory(case_id: str, memory_id: str) -> dict[str, object]:
    memory = store.get("case_memories", memory_id, CaseMemory)
    if not memory or memory.case_id != case_id:
        raise HTTPException(status_code=404, detail="Memory not found")
    store.delete("case_memories", memory_id)
    record("case.memory.deleted", "删除案件记忆", memory.content[:120], entity_type="case", entity_id=case_id)
    return {"ok": True}


@app.get("/api/cases/{case_id}/reply-jobs")
async def list_reply_jobs(case_id: str) -> list[LegalReplyJob]:
    if not store.get("cases", case_id, Case):
        raise HTTPException(status_code=404, detail="Case not found")
    return sorted(
        store.filter("reply_jobs", LegalReplyJob, case_id=case_id),
        key=lambda item: item.created_at,
        reverse=True,
    )


@app.post("/api/cases/{case_id}/reply-jobs")
async def create_reply_job(case_id: str, payload: ReplyJobCreateRequest) -> LegalReplyJob:
    case = store.get("cases", case_id, Case)
    if not case:
        raise HTTPException(status_code=404, detail="Case not found")
    memories = case_memories(case_id)
    messages = case_messages(case)
    mode_label = "短回复" if payload.mode == "short_reply" else "长回复"
    title = payload.title or f"{case.title}{mode_label}任务"
    summary = payload.case_summary.strip() or infer_case_summary(case, memories, messages)
    job = LegalReplyJob(
        case_id=case_id,
        mode=payload.mode,
        title=title,
        case_summary=summary,
        user_question=payload.user_question.strip(),
        assigned_agent_role=payload.assigned_agent_role,
        status="ready_for_review" if payload.mode == "short_reply" else "queued",
        draft_text=build_short_reply(case, payload, memories) if payload.mode == "short_reply" else "",
    )
    store.add("reply_jobs", job)
    record("reply_job.created", "创建回复工作流任务", title, entity_type="case", entity_id=case_id)
    return job


@app.post("/api/cases/{case_id}/reply-jobs/{job_id}/process")
async def process_reply_job(case_id: str, job_id: str) -> LegalReplyJob:
    case = store.get("cases", case_id, Case)
    if not case:
        raise HTTPException(status_code=404, detail="Case not found")
    job = store.get("reply_jobs", job_id, LegalReplyJob)
    if not job or job.case_id != case_id:
        raise HTTPException(status_code=404, detail="Reply job not found")
    if job.mode == "short_reply":
        job.status = "ready_for_review"
        job.module_state = "idle"
        job.updated_at = now_iso()
        if not job.draft_text:
            job.draft_text = build_short_reply(case, ReplyJobCreateRequest.model_validate(job.model_dump()), case_memories(case_id))
        store.update("reply_jobs", job)
        return job

    job.status = "reasoning"
    job.module_state = "busy"
    job.updated_at = now_iso()
    store.update("reply_jobs", job)

    runs = sorted(
        store.filter("legal_reasoning_runs", LegalReasoningRun, case_id=case_id),
        key=lambda item: item.created_at,
    )
    latest_run = runs[-1] if runs else None
    content = build_long_reply_content(
        case,
        job,
        case_memories(case_id),
        case_documents(case_id),
        latest_run,
    )
    document = create_text_document(
        case_id=case_id,
        title=f"{job.title} Word 输出",
        document_type="letter",
        content_text=content,
        source_filename=f"{job.title}.docx",
        change_summary="回复工作流自动生成",
        author_type="agent",
    )
    job.status = "completed"
    job.module_state = "idle"
    job.draft_text = content
    job.output_document_id = document.id
    job.updated_at = now_iso()
    store.update("reply_jobs", job)
    record("reply_job.completed", "长回复推理完成并生成文书", job.title, entity_type="case", entity_id=case_id)
    return job


@app.delete("/api/cases/{case_id}/reply-jobs/{job_id}")
async def delete_reply_job(case_id: str, job_id: str) -> dict[str, object]:
    job = store.get("reply_jobs", job_id, LegalReplyJob)
    if not job or job.case_id != case_id:
        raise HTTPException(status_code=404, detail="Reply job not found")
    store.delete("reply_jobs", job_id)
    record("reply_job.deleted", "删除回复工作流任务", job.title, entity_type="case", entity_id=case_id)
    return {"ok": True}


@app.post("/api/cases/{case_id}/follow-up-questions")
async def create_follow_up_question(case_id: str, payload: SendMessageRequest) -> FollowUpQuestion:
    if not store.get("cases", case_id, Case):
        raise HTTPException(status_code=404, detail="Case not found")
    question = FollowUpQuestion(case_id=case_id, content=payload.content)
    store.add("follow_up_questions", question)
    record("follow_up.created", "创建追问问题", question.content[:120], entity_type="case", entity_id=case_id)
    return question


@app.delete("/api/cases/{case_id}/follow-up-questions/{question_id}")
async def delete_follow_up_question(case_id: str, question_id: str) -> dict[str, object]:
    question = store.get("follow_up_questions", question_id, FollowUpQuestion)
    if not question or question.case_id != case_id:
        raise HTTPException(status_code=404, detail="Follow-up question not found")
    store.delete("follow_up_questions", question_id)
    record("follow_up.deleted", "删除追问问题", question.content[:120], entity_type="case", entity_id=case_id)
    return {"ok": True}


@app.post("/api/cases/{case_id}/follow-up-questions/{question_id}/send")
async def send_follow_up_question(
    case_id: str,
    question_id: str,
    payload: SendFollowUpRequest | None = None,
) -> dict[str, object]:
    case = store.get("cases", case_id, Case)
    if not case:
        raise HTTPException(status_code=404, detail="Case not found")
    if not case.conversation_ref:
        raise HTTPException(status_code=422, detail="Case is not bound to a WeChat conversation")
    conversation = store.get("wechat_conversations", case.conversation_ref, WechatConversation)
    if not conversation:
        raise HTTPException(status_code=404, detail="Conversation not found")
    question = store.get("follow_up_questions", question_id, FollowUpQuestion)
    if not question or question.case_id != case_id:
        raise HTTPException(status_code=404, detail="Follow-up question not found")

    content = (payload.content if payload and payload.content else question.content).strip()
    if not content:
        raise HTTPException(status_code=422, detail="Follow-up content is empty")

    if using_mock_wechat():
        new_msg = mock_wechat.append_message(
            conversation_id=conversation.id,
            sender="owner",
            content=content,
            source="manual",
        )
        message = WechatMessage.model_validate(new_msg)
        mock_wechat.sync_to_json_store(store)
        event_type = "follow_up.sent_mock"
        event_title = "通过演示模式发送追问"
    else:
        message = await OpenClawWechatAdapter(connection()).send_wechat_message(
            conversation_id=conversation.openclaw_conversation_id,
            content=content,
        )
        message.conversation_id = conversation.id
        store.add("wechat_messages", message)
        event_type = "follow_up.sent"
        event_title = "通过 OpenClaw 发送追问"

    conversation.last_message_at = message.created_at
    store.update("wechat_conversations", conversation)

    question.content = content
    question.status = "sent_via_openclaw" if message.status == "sent_via_openclaw" else "failed"
    question.sent_message_id = message.id
    question.failure_reason = (
        str(message.raw_payload.get("error")) if message.raw_payload and message.raw_payload.get("error") else None
    )
    question.updated_at = now_iso()
    store.update("follow_up_questions", question)
    record(
        event_type,
        event_title,
        content[:120],
        entity_type="case",
        entity_id=case_id,
    )
    return {"question": question, "message": message}


@app.delete("/api/reasoning/cases/{case_id}/runs/{run_id}")
async def delete_reasoning_run(case_id: str, run_id: str) -> dict[str, object]:
    run = store.get("legal_reasoning_runs", run_id, LegalReasoningRun)
    if not run or run.case_id != case_id:
        raise HTTPException(status_code=404, detail="Reasoning run not found")
    store.delete("legal_reasoning_runs", run_id)
    store.remove_where("follow_up_questions", lambda row: row.get("reasoning_run_id") == run_id)
    record("reasoning.deleted", "删除 AOE 推理图", run.input_summary[:120], entity_type="case", entity_id=case_id)
    return {"ok": True}


@app.get("/api/agents")
async def list_agents() -> list[LegalAgent]:
    return store.list("legal_agents", LegalAgent)


@app.get("/api/agents/architecture")
async def get_agent_architecture() -> AgentArchitecture:
    return build_agent_architecture()


@app.get("/api/documents")
async def list_documents(case_id: str | None = None) -> list[LegalDocument]:
    documents = store.list("legal_documents", LegalDocument)
    if case_id:
        documents = [item for item in documents if item.case_id == case_id]
    return sorted(documents, key=lambda item: item.updated_at, reverse=True)


@app.post("/api/documents")
async def create_document(payload: DocumentCreateRequest) -> LegalDocument:
    document = LegalDocument(
        case_id=payload.case_id,
        title=payload.title,
        document_type=payload.document_type,
    )
    revision = LegalDocumentRevision(
        document_id=document.id,
        version_number=1,
        content_text=payload.content_text,
        source_filename=payload.source_filename,
        change_summary=payload.change_summary,
    )
    create_default_branch_for_document(document, revision)
    store.add("legal_documents", document)
    store.add("legal_document_revisions", revision)
    record("document.created", "创建法律文件", document.title, entity_type="document", entity_id=document.id)
    return document


@app.post("/api/documents/upload")
async def upload_document(
    file: UploadFile = File(...),
    title: str | None = Form(default=None),
    case_id: str | None = Form(default=None),
    document_type: str = Form(default="other"),
    change_summary: str = Form(default="Initial upload"),
) -> LegalDocument:
    content_text = await read_legal_file_text(file)
    filename = file.filename or "uploaded document"
    if not content_text.strip():
        raise HTTPException(status_code=422, detail="Uploaded document has no readable text")
    normalized_type = document_type if document_type in {"contract", "letter", "pleading", "evidence", "other"} else "other"
    document = LegalDocument(
        case_id=case_id or None,
        title=title or Path(filename).stem,
        document_type=normalized_type,  # type: ignore[arg-type]
    )
    revision = LegalDocumentRevision(
        document_id=document.id,
        version_number=1,
        content_text=content_text,
        source_filename=filename,
        author_type="import",
        change_summary=change_summary,
    )
    create_default_branch_for_document(document, revision)
    store.add("legal_documents", document)
    store.add("legal_document_revisions", revision)
    record("document.uploaded", "上传法律文件", document.title, entity_type="document", entity_id=document.id)
    return document


@app.get("/api/documents/{document_id}")
async def get_document(document_id: str) -> dict[str, object]:
    document = store.get("legal_documents", document_id, LegalDocument)
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    revisions = store.filter("legal_document_revisions", LegalDocumentRevision, document_id=document_id)
    branches = document_branches(document_id)
    return {
        "document": document,
        "branches": branches,
        "revisions": sorted(revisions, key=lambda item: item.version_number),
    }


@app.get("/api/documents/{document_id}/tree")
async def get_document_tree(document_id: str) -> dict[str, object]:
    document = store.get("legal_documents", document_id, LegalDocument)
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    branches = document_branches(document_id)
    tree_branches = []
    for branch in branches:
        revisions = store.filter("legal_document_revisions", LegalDocumentRevision, document_id=document_id, branch_id=branch.id)
        revisions.sort(key=lambda item: item.version_number)
        tree_revisions = []
        for rev in revisions:
            label = f"v{rev.version_number} {rev.short_hash or ''} {rev.change_summary or '无版本说明'}"
            tree_revisions.append({
                "id": rev.id,
                "label": label,
                "version_number": rev.version_number,
                "short_hash": rev.short_hash,
                "parent_revision_id": rev.parent_revision_id,
                "change_summary": rev.change_summary,
                "source_filename": rev.source_filename,
                "created_at": rev.created_at,
            })
        tree_branches.append({
            **branch.model_dump(),
            "revisions": tree_revisions,
        })
    return {
        "document": document,
        "branches": tree_branches,
    }


@app.post("/api/documents/{document_id}/branches")
async def create_document_branch(document_id: str, payload: BranchCreateRequest) -> LegalDocumentBranch:
    document = store.get("legal_documents", document_id, LegalDocument)
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    name = payload.name.strip()
    if not name:
        raise HTTPException(status_code=422, detail="Branch name cannot be empty")
    import re
    if not re.match(r'^[\u4e00-\u9fa5a-zA-Z0-9_\-/]+$', name):
        raise HTTPException(status_code=422, detail="Branch name can only contain Chinese, English letters, digits, underscores, hyphens, and slashes")
    existing_branches = document_branches(document_id)
    if any(branch.name == name for branch in existing_branches):
        raise HTTPException(status_code=409, detail="Branch name already exists")
    base_revision = store.get("legal_document_revisions", payload.base_revision_id, LegalDocumentRevision)
    if not base_revision or base_revision.document_id != document_id:
        raise HTTPException(status_code=404, detail="Base revision not found")
    branch = LegalDocumentBranch(
        document_id=document_id,
        name=name,
        head_revision_id=base_revision.id,
        base_revision_id=base_revision.id,
        is_default=False,
    )
    store.add("legal_document_branches", branch)
    record("document.branch.created", "创建文件分支", f"{document.title} - {name}", entity_type="document", entity_id=document_id)
    return branch


@app.get("/api/documents/{document_id}/export.docx")
async def export_document_docx(document_id: str) -> StreamingResponse:
    document = store.get("legal_documents", document_id, LegalDocument)
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    revision = latest_document_revision(document_id)
    if not revision:
        raise HTTPException(status_code=404, detail="Document revision not found")
    stream = build_docx_stream(document.title, revision.content_text)
    filename = f"legal-document-{document.id}.docx"
    return StreamingResponse(
        stream,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# 7.9: POST /api/documents/{document_id}/branches/{branch_id}/revisions
@app.post("/api/documents/{document_id}/branches/{branch_id}/revisions")
async def create_branch_revision(
    document_id: str, branch_id: str, payload: BranchRevisionCreateRequest
) -> LegalDocumentRevision:
    document = store.get("legal_documents", document_id, LegalDocument)
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    branch = get_document_branch(document_id, branch_id)
    if not payload.content_text.strip():
        raise HTTPException(status_code=422, detail="Content text cannot be empty")
    revision = create_revision_on_branch(
        document=document,
        branch=branch,
        content_text=payload.content_text,
        source_filename=payload.source_filename,
        author_type=payload.author_type,
        change_summary=payload.change_summary,
    )
    record("document.revision.created", "向分支提交新版本", document.title, entity_type="document", entity_id=document.id)
    return revision


# 7.10: POST /api/documents/{document_id}/branches/{branch_id}/revisions/upload
@app.post("/api/documents/{document_id}/branches/{branch_id}/revisions/upload")
async def upload_branch_revision(
    document_id: str,
    branch_id: str,
    file: UploadFile = File(...),
    change_summary: str = Form(default="Uploaded revision"),
) -> LegalDocumentRevision:
    document = store.get("legal_documents", document_id, LegalDocument)
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    branch = get_document_branch(document_id, branch_id)
    content_text = await read_legal_file_text(file)
    if not content_text.strip():
        raise HTTPException(status_code=422, detail="Uploaded revision has no readable text")
    revision = create_revision_on_branch(
        document=document,
        branch=branch,
        content_text=content_text,
        source_filename=file.filename,
        author_type="import",
        change_summary=change_summary,
    )
    record("document.revision.uploaded", "向分支上传文件新版本", document.title, entity_type="document", entity_id=document.id)
    return revision


# 7.13: GET /api/documents/{document_id}/diff/export.docx
@app.get("/api/documents/{document_id}/diff/export.docx")
async def export_diff_docx(
    document_id: str,
    base_revision_id: str = Query(...),
    target_revision_id: str = Query(...),
) -> StreamingResponse:
    from docx.shared import RGBColor
    import re as _re

    document = store.get("legal_documents", document_id, LegalDocument)
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    base = store.get("legal_document_revisions", base_revision_id, LegalDocumentRevision)
    target = store.get("legal_document_revisions", target_revision_id, LegalDocumentRevision)
    if not base or base.document_id != document_id:
        raise HTTPException(status_code=404, detail="Base revision not found")
    if not target or target.document_id != document_id:
        raise HTTPException(status_code=404, detail="Target revision not found")

    segments = build_char_diff(base.content_text, target.content_text)
    paragraph_changes = build_paragraph_diff(base.content_text, target.content_text)
    risk_summary = summarize_legal_risks(base.content_text, target.content_text)

    doc = DocxDocument()
    doc.add_heading(f"{document.title} - 版本差异", level=1)

    base_label = f"v{base.version_number} {base.short_hash or ''} {base.change_summary or ''}".strip()
    target_label = f"v{target.version_number} {target.short_hash or ''} {target.change_summary or ''}".strip()
    doc.add_paragraph(f"基准版本：{base_label}")
    doc.add_paragraph(f"目标版本：{target_label}")

    doc.add_heading("逐字差异", level=2)
    for seg in segments:
        para = doc.paragraphs[-1] if doc.paragraphs else doc.add_paragraph()
        run = para.add_run(seg.text)
        if seg.op == "insert":
            run.font.color.rgb = RGBColor(22, 101, 52)  # green
        elif seg.op == "delete":
            run.font.color.rgb = RGBColor(153, 27, 27)  # red
            run.font.strike = True

    doc.add_heading("段落变化", level=2)
    for change in paragraph_changes:
        if change.op == "equal":
            continue
        para = doc.add_paragraph()
        if change.op == "insert":
            run = para.add_run(f"[新增] {change.target or ''}")
            run.font.color.rgb = RGBColor(22, 101, 52)
        elif change.op == "delete":
            run = para.add_run(f"[删除] {change.base or ''}")
            run.font.color.rgb = RGBColor(153, 27, 27)
            run.font.strike = True
        elif change.op == "replace":
            if change.base:
                run = para.add_run(f"原文：{change.base}")
                run.font.color.rgb = RGBColor(153, 27, 27)
                run.font.strike = True
            if change.target:
                para2 = doc.add_paragraph()
                run2 = para2.add_run(f"改为：{change.target}")
                run2.font.color.rgb = RGBColor(22, 101, 52)

    doc.add_heading("风险提示", level=2)
    for risk in risk_summary:
        doc.add_paragraph(risk, style="List Bullet")

    stream = BytesIO()
    doc.save(stream)
    stream.seek(0)
    filename = f"legal-document-diff-{document.id}.docx"
    return StreamingResponse(
        stream,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# 7.14: POST /api/documents/{document_id}/diff/analyze
@app.post("/api/documents/{document_id}/diff/analyze")
async def analyze_document_diff(
    document_id: str, payload: dict[str, str]
) -> LegalDocumentAnalysis:
    from urllib.request import Request, urlopen
    import json as _json

    document = store.get("legal_documents", document_id, LegalDocument)
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    base_revision_id = payload.get("base_revision_id", "")
    target_revision_id = payload.get("target_revision_id", "")
    base = store.get("legal_document_revisions", base_revision_id, LegalDocumentRevision)
    target = store.get("legal_document_revisions", target_revision_id, LegalDocumentRevision)
    if not base or base.document_id != document_id:
        raise HTTPException(status_code=404, detail="Base revision not found")
    if not target or target.document_id != document_id:
        raise HTTPException(status_code=404, detail="Target revision not found")

    # Check for cached analysis
    existing = store.filter(
        "legal_document_analyses",
        LegalDocumentAnalysis,
        document_id=document_id,
        base_revision_id=base.id,
        target_revision_id=target.id,
    )
    if existing:
        return existing[-1]

    # Generate diff data
    segments = build_char_diff(base.content_text, target.content_text)
    paragraph_changes = build_paragraph_diff(base.content_text, target.content_text)
    risk_summary = summarize_legal_risks(base.content_text, target.content_text)

    # Try LLM analysis
    api_key = os.getenv("LVZHIJIE_LLM_API_KEY") or os.getenv("OPENAI_API_KEY")
    base_url = (os.getenv("LVZHIJIE_LLM_BASE_URL") or os.getenv("OPENAI_BASE_URL") or "https://api.openai.com/v1").rstrip("/")
    model = os.getenv("LVZHIJIE_LLM_MODEL") or os.getenv("OPENAI_MODEL") or "gpt-4o-mini"

    analysis = None
    if api_key:
        try:
            changed_text = "\n".join(
                f"{base.content_text[max(0, i1 - 12):min(len(base.content_text), i2 + 12)]}\n"
                f"{target.content_text[max(0, j1 - 12):min(len(target.content_text), j2 + 12)]}"
                for op, i1, i2, j1, j2 in SequenceMatcher(
                    a=base.content_text, b=target.content_text
                ).get_opcodes()
                if op != "equal"
            )[:6000]

            paragraph_changes_text = "\n".join(
                f"[{change.op}] {change.base or ''} -> {change.target or ''}"
                for change in paragraph_changes
                if change.op != "equal"
            )[:3000]

            prompt = f"""请分析两个法律文件版本之间的差异。
只能输出 JSON 对象，不要输出 Markdown。
JSON 字段必须包含：
- risk_level: low / medium / high
- ambiguities: 字符串数组，列出可能存在歧义的变化
- stealth_changes: 字符串数组，列出疑似暗改、弱化责任、扩大免责、改变期限、改变争议解决方式的变化
- risk_points: 字符串数组，列出法律风险点
- suggestions: 字符串数组，列出修改或谈判建议
- manual_review_checklist: 字符串数组，列出人工必须复核的问题

文件标题：
{document.title}

基准版本：
v{base.version_number} {base.change_summary}

目标版本：
v{target.version_number} {target.change_summary}

段落变化：
{paragraph_changes_text}

逐字变化摘要：
{changed_text}"""

            request_body = _json.dumps({
                "model": model,
                "messages": [
                    {"role": "system", "content": "你是资深中国律师，专门审查合同版本差异、暗改风险和歧义条款。只能输出 JSON。"},
                    {"role": "user", "content": prompt},
                ],
                "temperature": 0.2,
            }).encode("utf-8")

            req = Request(
                f"{base_url}/chat/completions",
                data=request_body,
                headers={
                    "Content-Type": "application/json",
                    "Authorization": f"Bearer {api_key}",
                },
                method="POST",
            )
            with urlopen(req, timeout=30) as resp:
                result = _json.loads(resp.read().decode("utf-8"))
            content = result["choices"][0]["message"]["content"]
            if content.startswith("```"):
                content = _re.sub(r"^```(?:json)?\s*", "", content)
                content = _re.sub(r"\s*```$", "", content)
            parsed = _json.loads(content)
            analysis = LegalDocumentAnalysis(
                document_id=document_id,
                base_revision_id=base.id,
                target_revision_id=target.id,
                source="llm",
                risk_level=parsed.get("risk_level", "medium"),
                ambiguities=parsed.get("ambiguities", []),
                stealth_changes=parsed.get("stealth_changes", []),
                risk_points=parsed.get("risk_points", []),
                suggestions=parsed.get("suggestions", []),
                manual_review_checklist=parsed.get("manual_review_checklist", []),
            )
        except Exception:
            pass

    # Rule fallback
    if not analysis:
        analysis = LegalDocumentAnalysis(
            document_id=document_id,
            base_revision_id=base.id,
            target_revision_id=target.id,
            source="rule_fallback",
            risk_level="medium",
            ambiguities=["发现文本变化，请人工核对是否影响权利义务、履行期限或争议解决。"],
            stealth_changes=[],
            risk_points=risk_summary,
            suggestions=["建议逐条核对红色删除和绿色新增内容，确认是否改变双方实质权利义务。"],
            manual_review_checklist=[
                "核对付款金额、付款期限和付款条件是否变化。",
                "核对违约责任、赔偿上限和免责条款是否变化。",
                "核对解除条件、通知期限和争议解决条款是否变化。",
            ],
        )

    store.add("legal_document_analyses", analysis)
    return analysis


@app.delete("/api/documents/{document_id}")
async def delete_document(document_id: str) -> dict[str, object]:
    document = store.get("legal_documents", document_id, LegalDocument)
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    delete_document_rows(document_id)
    record("document.deleted", "删除法律文件", document.title, entity_type="document", entity_id=document_id)
    return {"ok": True}


@app.post("/api/documents/{document_id}/revisions")
async def create_revision(document_id: str, payload: RevisionCreateRequest) -> LegalDocumentRevision:
    document = store.get("legal_documents", document_id, LegalDocument)
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    branches = document_branches(document_id)
    default_branch = next((b for b in branches if b.is_default), None)
    if not default_branch:
        revisions = store.filter("legal_document_revisions", LegalDocumentRevision, document_id=document_id)
        if revisions:
            default_branch = create_default_branch_for_document(document, sorted(revisions, key=lambda r: r.version_number)[0])
        else:
            raise HTTPException(status_code=404, detail="No revisions found for this document")
    revision = create_revision_on_branch(
        document=document,
        branch=default_branch,
        content_text=payload.content_text,
        source_filename=payload.source_filename,
        author_type=payload.author_type,
        change_summary=payload.change_summary,
    )
    record("document.revision.created", "创建文件新版本", document.title, entity_type="document", entity_id=document.id)
    return revision


@app.delete("/api/documents/{document_id}/revisions/{revision_id}")
async def delete_revision(document_id: str, revision_id: str) -> dict[str, object]:
    document = store.get("legal_documents", document_id, LegalDocument)
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    revision = store.get("legal_document_revisions", revision_id, LegalDocumentRevision)
    if not revision or revision.document_id != document_id:
        raise HTTPException(status_code=404, detail="Revision not found")

    # 7.11: Prevent deleting branch head revisions
    branches = store.filter("legal_document_branches", LegalDocumentBranch, document_id=document_id)
    if any(branch.head_revision_id == revision_id for branch in branches):
        raise HTTPException(status_code=409, detail="Cannot delete branch head revision")

    # 7.11: Prevent deleting revisions with children
    child_revisions = store.filter(
        "legal_document_revisions", LegalDocumentRevision, document_id=document_id, parent_revision_id=revision_id
    )
    if child_revisions:
        raise HTTPException(status_code=409, detail="Cannot delete revision with children")

    store.delete("legal_document_revisions", revision_id)
    store.remove_where(
        "legal_document_diffs",
        lambda row: row.get("document_id") == document_id
        and (row.get("base_revision_id") == revision_id or row.get("target_revision_id") == revision_id),
    )
    store.remove_where(
        "legal_document_analyses",
        lambda row: row.get("document_id") == document_id
        and (row.get("base_revision_id") == revision_id or row.get("target_revision_id") == revision_id),
    )

    revisions = sorted(
        store.filter("legal_document_revisions", LegalDocumentRevision, document_id=document_id),
        key=lambda item: item.version_number,
    )
    if not revisions:
        delete_document_rows(document_id)
        record("document.deleted", "删除法律文件", document.title, entity_type="document", entity_id=document_id)
        return {"ok": True, "deleted_document": True}

    for index, item in enumerate(revisions, start=1):
        if item.version_number != index:
            item.version_number = index
            store.update("legal_document_revisions", item)
    # Update default branch head if needed
    default_branch = next((b for b in branches if b.is_default), None)
    if default_branch and default_branch.head_revision_id == revision_id:
        default_branch.head_revision_id = revisions[-1].id
        store.update("legal_document_branches", default_branch)
    document.current_revision_id = revisions[-1].id
    document.updated_at = now_iso()
    store.update("legal_documents", document)
    record("document.revision.deleted", "删除文件版本", document.title, entity_type="document", entity_id=document_id)
    return {"ok": True, "deleted_document": False}


@app.post("/api/documents/{document_id}/revisions/upload")
async def upload_revision(
    document_id: str,
    file: UploadFile = File(...),
    change_summary: str = Form(default="Uploaded revision"),
) -> LegalDocumentRevision:
    document = store.get("legal_documents", document_id, LegalDocument)
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    content_text = await read_legal_file_text(file)
    if not content_text.strip():
        raise HTTPException(status_code=422, detail="Uploaded revision has no readable text")
    branches = document_branches(document_id)
    default_branch = next((b for b in branches if b.is_default), None)
    if not default_branch:
        revisions = store.filter("legal_document_revisions", LegalDocumentRevision, document_id=document_id)
        if revisions:
            default_branch = create_default_branch_for_document(document, sorted(revisions, key=lambda r: r.version_number)[0])
        else:
            raise HTTPException(status_code=404, detail="No revisions found for this document")
    revision = create_revision_on_branch(
        document=document,
        branch=default_branch,
        content_text=content_text,
        source_filename=file.filename,
        author_type="import",
        change_summary=change_summary,
    )
    record("document.revision.uploaded", "上传文件新版本", document.title, entity_type="document", entity_id=document.id)
    return revision


@app.get("/api/documents/{document_id}/diff")
async def get_document_diff(
    document_id: str,
    base_revision_id: str | None = Query(default=None),
    target_revision_id: str | None = Query(default=None),
) -> LegalDocumentDiff:
    document = store.get("legal_documents", document_id, LegalDocument)
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    revisions = sorted(
        store.filter("legal_document_revisions", LegalDocumentRevision, document_id=document_id),
        key=lambda item: item.version_number,
    )
    if len(revisions) < 2 and not (base_revision_id and target_revision_id):
        raise HTTPException(status_code=422, detail="At least two revisions are required")
    base = store.get("legal_document_revisions", base_revision_id, LegalDocumentRevision) if base_revision_id else revisions[-2]
    target = (
        store.get("legal_document_revisions", target_revision_id, LegalDocumentRevision)
        if target_revision_id
        else revisions[-1]
    )
    if not base or base.document_id != document_id:
        raise HTTPException(status_code=404, detail="Base revision not found")
    if not target or target.document_id != document_id:
        raise HTTPException(status_code=404, detail="Target revision not found")
    diff = LegalDocumentDiff(
        document_id=document.id,
        base_revision_id=base.id,
        target_revision_id=target.id,
        segments=build_char_diff(base.content_text, target.content_text),
        paragraph_changes=build_paragraph_diff(base.content_text, target.content_text),
        risk_summary=summarize_legal_risks(base.content_text, target.content_text),
    )
    existing = store.filter(
        "legal_document_diffs",
        LegalDocumentDiff,
        document_id=document.id,
        base_revision_id=base.id,
        target_revision_id=target.id,
    )
    if existing:
        return existing[-1]
    store.add("legal_document_diffs", diff)
    return diff


@app.post("/api/reasoning/cases/{case_id}/generate")
async def generate_reasoning(case_id: str) -> LegalReasoningRun:
    case = store.get("cases", case_id, Case)
    if not case:
        raise HTTPException(status_code=404, detail="Case not found")
    memories = case_memories(case_id)
    messages = case_messages(case)
    documents = case_documents(case_id)
    summary = infer_case_summary(case, memories, messages)

    factual_memories = [memory for memory in memories if memory.kind in {"fact", "timeline", "evidence"}]
    uncertainty_memories = [memory for memory in memories if memory.kind == "uncertainty"]
    inbound_messages = [message for message in messages if message.direction == "inbound"]
    nodes: list[ReasoningNode] = []
    edges: list[ReasoningEdge] = []

    source_nodes: list[ReasoningNode] = []
    if factual_memories:
        for index, memory in enumerate(factual_memories[:4], start=1):
            node_type = "Timeline" if memory.kind == "timeline" else "Evidence" if memory.kind == "evidence" else "Fact"
            node = ReasoningNode(
                node_type=node_type,
                label=f"{'时间线' if node_type == 'Timeline' else '证据' if node_type == 'Evidence' else '事实'} {index}",
                content=memory.content,
                confidence=memory.confidence,
                source_refs=[memory.source_ref] if memory.source_ref else [],
            )
            nodes.append(node)
            source_nodes.append(node)
    elif inbound_messages:
        node = ReasoningNode(
            node_type="Fact",
            label="咨询事实",
            content=first_sentence(inbound_messages[0].content, summary),
            confidence=0.58,
            source_refs=[inbound_messages[0].id],
        )
        nodes.append(node)
        source_nodes.append(node)
    else:
        node = ReasoningNode(node_type="Fact", label="案件摘要", content=summary, confidence=0.5)
        nodes.append(node)
        source_nodes.append(node)

    if documents:
        document_node = ReasoningNode(
            node_type="Evidence",
            label="案件文件",
            content="；".join(f"{document.title}（{document.document_type}）" for document in documents[:4]),
            confidence=0.7,
            source_refs=[document.id for document in documents[:4]],
        )
        nodes.append(document_node)
        source_nodes.append(document_node)

    issue = ReasoningNode(
        node_type="Issue",
        label=f"{case_type_label(case.case_type)}争议焦点",
        content=f"需要判断本案在{case_type_label(case.case_type)}框架下的请求基础、关键事实和证明责任。",
        confidence=0.68,
    )
    rule = ReasoningNode(
        node_type="Rule",
        label="法律要件",
        content=case_rule_hint(case.case_type),
        confidence=0.64,
    )
    analysis = ReasoningNode(
        node_type="Analysis",
        label="阶段分析",
        content=f"当前摘要：{summary} 现阶段应先核对事实与证据，再形成可对外使用的法律意见。",
        confidence=0.66,
    )
    conclusion = ReasoningNode(
        node_type="Conclusion",
        label="阶段结论",
        content="可以形成内部阶段性判断；如关键证据未补齐，应先追问并人工复核后再输出正式结论。",
        confidence=0.6,
    )
    nodes.extend([issue, rule, analysis, conclusion])

    for source in source_nodes:
        edges.append(ReasoningEdge(source=source.id, target=issue.id, relation_type="supports"))
    edges.extend(
        [
            ReasoningEdge(source=issue.id, target=rule.id, relation_type="requires"),
            ReasoningEdge(source=rule.id, target=analysis.id, relation_type="supports"),
            ReasoningEdge(source=analysis.id, target=conclusion.id, relation_type="leads_to"),
        ]
    )

    missing_points = [memory.content for memory in uncertainty_memories]
    if not documents:
        missing_points.append("尚未上传可核对的合同、证据或案件材料。")
    if not any(memory.kind == "timeline" for memory in memories):
        missing_points.append("尚未形成完整时间线。")
    if not missing_points:
        missing_points.append("需由人工复核事实来源与对外表述边界。")

    uncertainty_node = ReasoningNode(
        node_type="Uncertainty",
        label="推理暂停点",
        content="；".join(missing_points[:3]),
        confidence=0.88,
    )
    nodes.append(uncertainty_node)
    edges.append(ReasoningEdge(source=analysis.id, target=uncertainty_node.id, relation_type="uncertain_about"))

    follow_up_questions = [
        "请补充关键时间节点，并按发生顺序说明每一步沟通或履行情况。",
        "请列明目前已有证据材料，包括合同、聊天记录、转账/工资流水、通知文件或其他书面材料。",
        "请确认您的目标是协商解决、发送函件、投诉举报，还是准备仲裁/诉讼。",
    ]
    if case.case_type == "labor":
        follow_up_questions.insert(1, "请补充入职时间、劳动合同签订情况、工资标准、欠薪月份和离职沟通证据。")
    elif case.case_type == "criminal":
        follow_up_questions.insert(1, "请补充涉案行为发生时间、地点、参与人员、目前程序阶段和已取得的法律文书。")
    elif case.case_type == "contract":
        follow_up_questions.insert(1, "请补充合同签订时间、履行节点、违约事实、催告记录和争议条款。")

    question_node = ReasoningNode(
        node_type="Question",
        label="下一轮追问",
        content="；".join(follow_up_questions[:4]),
        confidence=0.92,
    )
    nodes.append(question_node)
    edges.append(ReasoningEdge(source=uncertainty_node.id, target=question_node.id, relation_type="asks"))

    needs_evidence = bool(missing_points and missing_points[0] != "需由人工复核事实来源与对外表述边界。")
    run = LegalReasoningRun(
        case_id=case_id,
        status="needs_evidence" if needs_evidence else "ready_for_review",
        input_summary=summary,
        nodes=nodes,
        edges=edges,
        follow_up_questions=follow_up_questions,
        blocked_reason="；".join(missing_points[:3]) if needs_evidence else "",
        review_focus=["事实来源", "证据完整性", "法律依据", "对外表述边界"],
        output_summary=f"已生成{len(nodes)}个节点、{len(edges)}条关系，建议先处理追问后再输出正式意见。",
    )
    store.add("legal_reasoning_runs", run)
    for content in follow_up_questions:
        store.add(
            "follow_up_questions",
            FollowUpQuestion(case_id=case_id, reasoning_run_id=run.id, content=content),
        )
    record("reasoning.generated", "生成 AOE 推理图", case.title, entity_type="case", entity_id=case_id)
    return run


@app.get("/api/mock-wechat/conversations")
async def list_mock_conversations() -> list[dict[str, object]]:
    return mock_wechat.list_conversations()


@app.post("/api/mock-wechat/conversations")
async def create_mock_conversation(payload: MockConversationCreateRequest) -> dict[str, object]:
    return mock_wechat.create_conversation(
        display_name=payload.display_name,
        remark=payload.remark,
        avatar_url=payload.avatar_url,
    )


@app.put("/api/mock-wechat/conversations/{conversation_id}")
async def update_mock_conversation(
    conversation_id: str, payload: MockConversationUpdateRequest
) -> dict[str, object]:
    raw = {k: v for k, v in payload.model_dump().items() if v is not None}
    contact_fields = {}
    update_data: dict[str, object] = {}
    for key, value in raw.items():
        if key in {"display_name", "remark", "avatar_url"}:
            contact_fields[key] = value
        else:
            update_data[key] = value
    if contact_fields:
        update_data["contact"] = contact_fields
    result = mock_wechat.update_conversation(conversation_id, update_data)
    if result is None:
        raise HTTPException(status_code=404, detail="Mock conversation not found")
    return result


@app.delete("/api/mock-wechat/conversations/{conversation_id}")
async def delete_mock_conversation(conversation_id: str) -> dict[str, object]:
    conversation = store.get("wechat_conversations", conversation_id, WechatConversation)
    if not mock_wechat.delete_conversation(conversation_id):
        raise HTTPException(status_code=404, detail="Mock conversation not found")
    if conversation:
        delete_wechat_conversation_rows(conversation)
    return {"ok": True}


@app.get("/api/mock-wechat/conversations/{conversation_id}/messages")
async def list_mock_messages(conversation_id: str) -> list[dict[str, object]]:
    return mock_wechat.list_messages(conversation_id)


@app.post("/api/mock-wechat/conversations/{conversation_id}/messages")
async def create_mock_message(
    conversation_id: str,
    sender: str = Form(...),
    content: str = Form(default=""),
    files: list[UploadFile] | None = File(default=None),
) -> dict[str, object]:
    if sender not in {"wechat_user", "owner"}:
        raise HTTPException(status_code=422, detail="sender must be wechat_user or owner")
    uploads = files or []
    if not content.strip() and not uploads:
        raise HTTPException(status_code=422, detail="Message content or files are required")
    attachments = []
    for upload in uploads:
        attachment = await mock_wechat.save_upload(upload)
        attachments.append(attachment)
    new_msg = mock_wechat.append_message(
        conversation_id=conversation_id,
        sender=sender,
        content=content,
        attachments=attachments,
        source="mock",
    )
    mock_wechat.sync_to_json_store(store)
    return new_msg


@app.delete("/api/mock-wechat/conversations/{conversation_id}/messages/{message_id}")
async def delete_mock_message(conversation_id: str, message_id: str) -> dict[str, object]:
    if not mock_wechat.delete_message(conversation_id, message_id):
        raise HTTPException(status_code=404, detail="Mock message not found")
    return {"ok": True}
